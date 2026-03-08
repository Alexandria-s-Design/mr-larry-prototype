"""
Generate the Life Dreams Wealth curriculum proposal for Mr. Larry Wimsatt.
V2 — Standards-aligned, game-forward, integrated life experience philosophy.
No pricing or timeline. Focus on curriculum, standards, and game design.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

REPO = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(REPO, "Life_Dreams_Wealth_Curriculum_Proposal.docx")

NAVY = RGBColor(0x0F, 0x0F, 0x1A)
GOLD = RGBColor(0xC9, 0xA8, 0x4C)
DARK_GOLD = RGBColor(0xA0, 0x85, 0x3D)
TEAL = RGBColor(0x2A, 0x9D, 0x8F)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK_TEXT = RGBColor(0x1A, 0x1A, 0x2E)
MID_GRAY = RGBColor(0x55, 0x55, 0x55)
CORAL = RGBColor(0xE0, 0x6C, 0x5A)
PURPLE = RGBColor(0x6B, 0x5B, 0x95)

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = DARK_TEXT
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15


def shade(cell, color_hex):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'))


def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    colors = {1: NAVY, 2: DARK_GOLD, 3: TEAL}
    sizes = {1: Pt(22), 2: Pt(16), 3: Pt(13)}
    for r in h.runs:
        r.font.color.rgb = colors.get(level, NAVY)
        r.font.size = sizes.get(level, Pt(22))
        r.font.name = 'Calibri'
    return h


def body(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(11)
    r.font.name = 'Calibri'
    r.font.color.rgb = DARK_TEXT
    return p


def bullet(text, bp=""):
    p = doc.add_paragraph(style='List Bullet')
    if bp:
        r = p.add_run(bp)
        r.bold = True
        r.font.size = Pt(11)
        r.font.name = 'Calibri'
        r.font.color.rgb = DARK_TEXT
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.name = 'Calibri'
    r.font.color.rgb = DARK_TEXT
    return p


def table(headers, rows, widths=None, hdr_color="0F0F1A"):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = WHITE
        r.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade(c, hdr_color)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            c.text = ""
            p = c.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            r.font.name = 'Calibri'
            r.font.color.rgb = DARK_TEXT
            if ri % 2 == 0:
                shade(c, "F5F5F0")
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t


# =====================================================================
# COVER PAGE
# =====================================================================
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("LIFE DREAMS WEALTH")
r.font.size = Pt(36)
r.font.color.rgb = NAVY
r.bold = True
r.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("An Interactive Financial Literacy Game & Curriculum")
r.font.size = Pt(18)
r.font.color.rgb = DARK_GOLD
r.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Standards-Aligned. Family-Centered. Game-Driven.")
r.font.size = Pt(14)
r.font.color.rgb = TEAL
r.font.name = 'Calibri'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("_" * 60)
r.font.color.rgb = TEAL
r.font.size = Pt(10)

doc.add_paragraph()

for txt in [
    "Prepared for Mr. Lawrence (Larry) Wimsatt",
    "Life Dreams Wealth | Long Beach, CA",
    "",
    "Prepared by Alexandria's Design LLC",
    "Dr. Charles Martin | Moreno Valley, CA | 310-709-4893",
    "",
    "March 2026 | Curriculum & Game Design Proposal",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt)
    r.font.size = Pt(11)
    r.font.color.rgb = DARK_TEXT
    r.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(1)

doc.add_page_break()

# =====================================================================
# TABLE OF CONTENTS
# =====================================================================
heading("Table of Contents", 1)
for item in [
    "1. Executive Summary",
    "2. Mr. Larry's Vision: Life Experiences as Financial Education",
    "3. The Life-Dreams-Wealth Framework",
    "4. National Standards Alignment",
    "5. Curriculum Standards Crosswalk",
    "6. Target Audience & Developmental Design",
    "7. Platform Architecture: The Dual-Portal Family Game",
    "8. Game Design Philosophy & Mechanics",
    "9. The 3D World: Navigating Financial Life",
    "10. The Game Engine: Phenomenon-Based Simulation Lessons",
    "11. Complete Simulation Library (K-12)",
    "12. Family Gameplay: The Parent-Child Connection",
    "13. Assessment Through Play",
    "14. Prior Work & Existing Assets",
    "15. Next Steps",
]:
    p = doc.add_paragraph()
    r = p.add_run(item)
    r.font.size = Pt(12)
    r.font.name = 'Calibri'
    r.font.color.rgb = DARK_TEXT
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# =====================================================================
# 1. EXECUTIVE SUMMARY
# =====================================================================
heading("1. Executive Summary", 1)

body(
    "Life Dreams Wealth (LDW) is a family financial literacy game and curriculum that transforms how "
    "families learn about money by grounding every lesson in real life experiences. It is not a textbook "
    "with quizzes. It is not a lecture series. It is an interactive, narrative-driven game world where "
    "players navigate the same financial decisions they will face in life — from a child splitting allowance "
    "into Save, Spend, and Share jars, to a high schooler surviving their first 90 days after graduation."
)

body(
    "Created by Mr. Lawrence Wimsatt, LDW is built on a core insight that most financial education misses: "
    "money decisions are never just about money. They are about identity, family, community, access, and "
    "the compounding effect of life circumstances. A grocery store trip teaches budgeting, nutrition, and "
    "time management simultaneously. A broken-down car teaches emergency funds, insurance, credit, and "
    "the hidden cost of living in a neighborhood without a mechanic. Every financial moment is a life moment."
)

body(
    "This proposal details the full curriculum — 20 phenomenon-based simulation lessons across K-12, "
    "aligned to national financial literacy standards (Jump$tart Coalition, Council for Economic Education, "
    "and state frameworks), delivered through an immersive game platform with dual Parent/Child portals."
)

doc.add_page_break()

# =====================================================================
# 2. MR. LARRY'S VISION
# =====================================================================
heading("2. Mr. Larry's Vision: Life Experiences as Financial Education", 1)

body(
    '"Breaking Financial Barriers: Empowering Diverse Communities"',
    bold=True,
)

body(
    "Mr. Wimsatt's philosophy begins where most curricula end. Traditional financial literacy teaches "
    "concepts in isolation: here is a budget worksheet, here is compound interest, here is a credit score. "
    "But that's not how people experience money. People experience money through life — through the "
    "grocery store, the car breaking down, the first apartment, the birthday party they can't quite afford, "
    "the conversation at the dinner table about whether to fix the roof or save for vacation."
)

heading("The Integrated Life Experience Model", 3)
body(
    "In LDW, every lesson begins with a life experience — not a definition. Students don't learn what a "
    "budget is and then apply it. They are dropped into a scenario where they NEED a budget to survive, "
    "and they discover the concept through play. This is phenomenon-based learning applied to financial "
    "literacy: start with a real situation, let the learner struggle with it, then name the concept."
)

doc.add_paragraph()

bullet(" Going grocery shopping with $40 teaches budgeting, comparison shopping, nutrition, "
       "unit pricing, opportunity cost, and impulse control — all in one experience.", bp="The Grocery Store: ")
bullet(" Planning a birthday on $20 teaches wants vs. needs, prioritization, and the discovery that "
       "free things (musical chairs) can be just as fun as expensive things (bounce houses).", bp="The Birthday Party: ")
bullet(" Getting a paycheck and watching deductions shrink it teaches taxes, net vs. gross, fixed "
       "expenses, and the emotional shock that prepares kids for adult financial reality.", bp="The First Paycheck: ")
bullet(" Two families with the same income but different neighborhoods — one near a credit union, "
       "one near payday lenders — reveals how access and systems shape outcomes.", bp="The Neighborhood: ")
bullet(" Twins choose different paths: college vs. trade. Same graduation day, wildly different "
       "financial trajectories over 40 years. Neither is wrong — but the math tells different stories.", bp="The Career Fork: ")
bullet(" A denied home loan in 1960 cascades across three generations into a $450K wealth gap "
       "by 2026. Same work ethic. Different rules. Historical context changes everything.", bp="The Generational Story: ")

doc.add_paragraph()

heading("The Iceberg Metaphor", 3)
body(
    "A recurring visual throughout LDW is the iceberg. What the world sees — someone's car, clothes, "
    "apartment — is just the tip. Beneath the surface: savings habits, credit history, investment knowledge, "
    "generational wealth or generational debt, access to resources, and the mindset driving every decision. "
    "LDW teaches families to build from the bottom up. The game world literally uses an iceberg as a "
    "progress visualization — as players learn, more of the iceberg rises above water."
)

heading("Family as the Unit of Play", 3)
body(
    "Most financial games target individuals. LDW treats the family as the player. Parents have their own "
    "portal with adult-level challenges (managing payroll, reviewing household budgets, tracking their "
    "child's progress). Children have their own portal with age-appropriate gameplay (earning virtual money "
    "through chores, managing savings, running simulations). But the game connects them — shared family "
    "goals, collaborative challenges, and dinner-table conversation starters that make financial learning "
    "a natural part of family life."
)

doc.add_page_break()

# =====================================================================
# 3. THE FRAMEWORK
# =====================================================================
heading("3. The Life-Dreams-Wealth Framework", 1)

body(
    "Every experience in the game maps to one of three phases. This is not a linear progression — it is "
    "a cycle. Players revisit each phase with deeper understanding as they grow."
)

doc.add_paragraph()

table(
    ["Phase", "Core Question", "Life Experiences", "Financial Concepts", "Game Mechanic"],
    [
        ["LIFE\n(Where You Are)", "What is my financial reality right now?",
         "Grocery shopping, getting a paycheck, paying bills, comparing neighborhoods, understanding taxes",
         "Budgeting, income, expenses, net vs. gross, taxes, needs vs. wants, opportunity cost",
         "Exploration & Discovery — open-world navigation, uncovering hidden information, reality check meters"],
        ["DREAMS\n(Where You're Going)", "What do I want, and how do I get there?",
         "Planning a birthday, saving for a toy, fundraising for a class trip, starting a business, choosing a career",
         "Goal-setting (short/intermediate/long), saving strategies, entrepreneurship, career ROI, planning",
         "Quest Systems — goal trackers, milestone rewards, branching story paths, countdown timers"],
        ["WEALTH\n(Building Resources)", "How do I build and protect what I have?",
         "Opening a savings account, investing, buying insurance, understanding credit, building generational legacy",
         "Compound interest, investing, insurance, credit scores, generational wealth, estate basics",
         "Strategy & Building — portfolio builders, compound growth visualizers, legacy planning, shield mechanics"],
    ],
    widths=[0.9, 1.1, 1.5, 1.5, 1.5],
)

doc.add_paragraph()

heading("Goal Types Within DREAMS", 3)
body("Mr. Larry identifies three goal horizons that map to children's lived experience of time:")
bullet(" 4 weeks — achievable within one topic cycle. 'Save $12 for art supplies.'", bp="Short-Term: ")
bullet(" 10-12 weeks — requires sustained effort across multiple topics. 'Save $50 for a video game.'", bp="Intermediate: ")
bullet(" 26 weeks — a half-year vision. 'Save $150 for a family experience.'", bp="Long-Term: ")
body(
    "The game visualizes these as quest lines of different lengths. Short quests complete within one world. "
    "Intermediate quests span multiple worlds. Long quests run as background missions across the entire "
    "game experience, with periodic check-ins and celebration moments."
)

heading("The Action Plan: Fitting Dreams Into Current Life", 3)
body(
    "Every DREAMS-phase lesson ends with the same question: 'How does this goal fit into your current plan?' "
    "This is Mr. Larry's key insight — goals without grounding in reality are fantasies. The game requires "
    "players to map their dreams against their LIFE-phase reality (actual income, actual expenses, actual "
    "constraints) before the quest can begin. This teaches planning discipline through gameplay, not lectures."
)

doc.add_page_break()

# =====================================================================
# 4. NATIONAL STANDARDS ALIGNMENT
# =====================================================================
heading("4. National Standards Alignment", 1)

body(
    "LDW is designed to meet or exceed every major national standard for K-12 financial literacy. "
    "This ensures the curriculum is not only engaging but adoptable by schools, districts, after-school "
    "programs, and community organizations that require standards-aligned content."
)

heading("Jump$tart Coalition National Standards for K-12 Personal Financial Literacy", 2)
body(
    "The Jump$tart Coalition is the leading national authority on K-12 financial literacy standards, "
    "endorsed by 150+ organizations including the Federal Reserve, FDIC, and National Endowment for "
    "Financial Education. LDW maps to all six Jump$tart domains:"
)

doc.add_paragraph()
table(
    ["Domain", "Code", "Jump$tart Description", "How LDW Addresses It"],
    [
        ["Spending & Saving", "SS",
         "Apply strategies for managing personal finances including spending, saving, and budgeting",
         "Grocery Store sim, Birthday Budget, Paycheck Payday, Three Piggy Banks, Apartment Hunt — every LIFE-phase lesson requires active budgeting through gameplay"],
        ["Credit & Debt", "CD",
         "Understand the purpose and responsible use of credit and the consequences of excessive debt",
         "Credit Score Mystery (6-8), Generational Wealth Gap, Community Resource Map — students see credit as a system, not just a number"],
        ["Employment & Income", "EI",
         "Use career planning to develop personal income potential and understand compensation",
         "Lemonade Stand, Neighborhood Business, Paycheck Payday, College vs. Career, Life After Graduation — earning is explored through entrepreneurship AND employment"],
        ["Investing", "INV",
         "Evaluate investment alternatives and apply strategies to achieve financial goals",
         "Savings Snowball, Investment Challenge, Building Generational Wealth — compound growth is visualized through animated game mechanics"],
        ["Risk Management & Insurance", "RM",
         "Apply strategies to manage risk and insure against financial loss",
         "Insurance Gamble, Life After Graduation emergency events, Apartment Hunt random life events — risk is experienced as gameplay surprise mechanics"],
        ["Financial Decision Making", "FDM",
         "Apply reliable information and systematic decision-making to personal financial decisions",
         "EVERY simulation in LDW — branching decisions with visible consequences are the core game mechanic; every choice has a traceable outcome"],
    ],
    widths=[1.0, 0.4, 2.0, 3.1],
)

doc.add_page_break()

heading("Council for Economic Education (CEE) National Standards", 2)
body(
    "The CEE standards focus on economic reasoning and personal finance. LDW aligns to the six CEE "
    "Personal Finance standards:"
)

doc.add_paragraph()
table(
    ["CEE Standard", "Description", "LDW Alignment"],
    [
        ["Earning Income", "People earn income by working, and income varies by job, education, skills, and market conditions",
         "Paycheck Payday, Lemonade Stand, Neighborhood Business, College vs. Career — students experience the full spectrum from chores to careers"],
        ["Buying Goods & Services", "People purchase goods and services to satisfy wants, using comparison shopping and budgeting",
         "HealthMart Foods (flagship sim), Birthday Party Budget, Apartment Hunt — comparison shopping is the core mechanic"],
        ["Saving", "Saving is choosing to set aside income for future use; interest compounds over time",
         "Three Piggy Banks, Savings Snowball, Class Trip Fund — saving is gamified with visual growth trackers and 'interest fairy' animations"],
        ["Using Credit", "Credit allows people to buy now and pay later, but carries costs and risks",
         "Credit Score Mystery, Shark Tank Jr. funding choices, Generational Wealth Gap redlining history"],
        ["Financial Investing", "Financial investment earns income through interest, dividends, and capital gains",
         "Investment Challenge (30-year sim with real historical data), Building Generational Wealth (3-generation portfolio)"],
        ["Protecting & Insuring", "People use insurance and risk management to protect against financial loss",
         "Insurance Gamble (5-year dual-household simulation), emergency events woven into multiple sims"],
    ],
    widths=[1.3, 2.2, 3.0],
)

doc.add_page_break()

heading("State Standards Alignment", 2)
body(
    "As of 2026, 35 states require or encourage financial literacy education. LDW is designed to align "
    "with the most widely adopted state frameworks:"
)

doc.add_paragraph()
table(
    ["State Framework", "Key Requirements", "LDW Coverage"],
    [
        ["California (AB 984, 2026)", "Financial literacy required for graduation starting 2027. Must cover budgeting, credit, saving, investing, insurance, taxes, and identity theft.",
         "Full coverage across all 20 simulations. Taxes (Tax Return sim), identity theft awareness integrated into Credit Score Mystery."],
        ["Florida (HB 1303)", "Standalone semester course required for high school graduation. Covers earning, spending, saving, investing, credit, debt.",
         "LDW's 9-12 simulations cover all required topics. Can serve as the curriculum backbone for the mandatory course."],
        ["Texas (SB 1590)", "Personal financial literacy required in middle and high school. Includes budgeting, credit, insurance, investing.",
         "6-8 and 9-12 grade bands provide comprehensive coverage. Standards crosswalk available for TEKS alignment."],
        ["Virginia (New 2025)", "K-12 integrated financial literacy across subjects. Emphasis on experiential learning.",
         "LDW's phenomenon-based approach IS experiential learning. K-2 through 9-12 coverage matches the integrated mandate."],
        ["Common Core Math", "Multiple standards reference financial applications: ratios, percentages, statistics, functions.",
         "Every simulation uses math in context — unit pricing, compound interest calculations, percentage-based budgets, probability."],
    ],
    widths=[1.3, 2.5, 2.7],
)

doc.add_paragraph()
body(
    "LDW is positioned to be a turnkey solution for states adopting new financial literacy requirements. "
    "The game format makes compliance engaging rather than burdensome for students and teachers.",
    italic=True,
)

doc.add_page_break()

# =====================================================================
# 5. CURRICULUM CROSSWALK
# =====================================================================
heading("5. Curriculum Standards Crosswalk", 1)

body(
    "The following crosswalk maps every LDW simulation to the specific standards it addresses. "
    "This is the document schools and organizations need to adopt LDW as an approved curriculum."
)

doc.add_paragraph()
table(
    ["Sim #", "Simulation Title", "Grade", "Phase", "SS", "CD", "EI", "INV", "RM", "FDM"],
    [
        ["K2-01", "The Birthday Party Budget", "K-2", "LIFE", "X", "", "", "", "", "X"],
        ["K2-02", "The Lemonade Stand", "K-2", "DREAMS", "X", "", "X", "", "", "X"],
        ["K2-03", "The Coin Quest", "K-2", "LIFE", "X", "", "", "", "", "X"],
        ["K2-04", "The Three Piggy Banks", "K-2", "DREAMS", "X", "", "", "", "", "X"],
        ["35-01", "HealthMart Foods", "3-5", "LIFE", "X", "", "", "", "", "X"],
        ["35-02", "The Class Trip Fund", "3-5", "DREAMS", "X", "", "X", "", "", "X"],
        ["35-03", "Paycheck Payday", "3-5", "LIFE", "X", "", "X", "", "", ""],
        ["35-04", "The Savings Snowball", "3-5", "WEALTH", "X", "", "", "X", "", "X"],
        ["35-05", "The Neighborhood Business", "3-5", "DREAMS", "", "", "X", "", "", "X"],
        ["35-06", "Community Resource Map", "3-5", "WEALTH", "X", "X", "", "", "", "X"],
        ["68-01", "Credit Score Mystery", "6-8", "LIFE", "", "X", "", "", "", "X"],
        ["68-02", "The Investment Challenge", "6-8", "WEALTH", "", "", "", "X", "", "X"],
        ["68-03", "The Apartment Hunt", "6-8", "LIFE", "X", "", "", "", "X", "X"],
        ["68-04", "Shark Tank Jr.", "6-8", "DREAMS", "", "X", "X", "", "", "X"],
        ["68-05", "The Insurance Gamble", "6-8", "WEALTH", "", "", "", "", "X", "X"],
        ["68-06", "Generational Wealth Gap", "6-8", "WEALTH", "", "X", "", "X", "", "X"],
        ["912-01", "Life After Graduation", "9-12", "LIFE", "X", "X", "X", "", "X", "X"],
        ["912-02", "The Tax Return", "9-12", "LIFE", "", "", "X", "", "", "X"],
        ["912-03", "College vs. Career", "9-12", "DREAMS", "", "X", "X", "", "", "X"],
        ["912-04", "Generational Wealth", "9-12", "WEALTH", "", "", "", "X", "", "X"],
    ],
    widths=[0.5, 1.5, 0.5, 0.55, 0.3, 0.3, 0.3, 0.3, 0.3, 0.35],
)

doc.add_paragraph()
body("SS = Spending & Saving | CD = Credit & Debt | EI = Employment & Income", italic=True)
body("INV = Investing | RM = Risk Management | FDM = Financial Decision Making", italic=True)
body(
    "Coverage: 100% of Jump$tart domains addressed. FDM (Financial Decision Making) is embedded in every "
    "simulation — it is the core gameplay mechanic. The remaining 5 domains are distributed across grade "
    "bands with increasing sophistication."
)

doc.add_page_break()

# =====================================================================
# 6. TARGET AUDIENCE
# =====================================================================
heading("6. Target Audience & Developmental Design", 1)

body(
    "LDW serves four grade bands, each designed for the cognitive, emotional, and social development "
    "stage of the learner. The same Life-Dreams-Wealth framework scales across all ages — a kindergartner "
    "and a high schooler both explore 'where am I now?' but with vastly different depth."
)

doc.add_paragraph()
table(
    ["Grade Band", "Ages", "Cognitive Stage", "Game Style", "Life Experiences Used"],
    [
        ["K-2", "5-8", "Concrete, sensory, narrative-driven. Understands 'mine' vs. 'ours.' Beginning to grasp cause and effect.",
         "Guided story games with 2-way choices. Animated characters. Audio narration. Drag-and-drop interactions. Bright, warm environments.",
         "Birthday parties, lemonade stands, coin jars, allowance splitting, toy stores"],
        ["3-5", "8-11", "Operational thinking. Can compare, calculate, and plan short-term. Understands fairness and rules.",
         "Interactive branching sims with 3-way choices. Real math (unit pricing, percentages). Visual dashboards. Achievement badges.",
         "Grocery shopping, class fundraisers, paychecks, savings accounts, starting a neighborhood business, exploring the community"],
        ["6-8", "11-14", "Abstract reasoning emerging. Can handle percentages, projections, and systemic thinking. Social identity forming.",
         "Complex multi-path sims. Spreadsheet-style tools. Real-world data. Multiplayer elements. 'What-if' scenario engines. Social comparison mechanics.",
         "Renting an apartment, building credit, investing, starting a real business, buying insurance, understanding generational inequity"],
        ["9-12", "14-18", "Full abstract reasoning. Can model long-term outcomes. Identity and values solidifying. Approaching real financial decisions.",
         "Full financial modeling. Portfolio management. Tax simulation. 40-year life projectors. Real job/school data. Mentor NPC (Mr. Larry) as career guide.",
         "Surviving post-graduation, filing taxes, choosing college vs. career, building generational wealth, creating a family legacy plan"],
    ],
    widths=[0.5, 0.4, 1.5, 1.5, 2.6],
)

doc.add_paragraph()

heading("Cultural Responsiveness", 3)
body(
    "LDW is designed with intentional cultural responsiveness for African American and Hispanic American "
    "communities — populations disproportionately affected by the wealth gap, predatory lending, and "
    "lack of access to financial education. This is not about creating 'diversity content' — it is about "
    "designing scenarios that reflect the real financial landscape these families navigate. A simulation "
    "about choosing between a credit union and a check-cashing store isn't theoretical for many families — "
    "it is Tuesday. LDW meets families where they are."
)

doc.add_page_break()

# =====================================================================
# 7. PLATFORM ARCHITECTURE
# =====================================================================
heading("7. Platform Architecture: The Dual-Portal Family Game", 1)

heading("The Family Login", 3)
body(
    "Players open the app and enter their Family Name. Two buttons appear: PARENT and KIDS. Each opens "
    "a different game portal. An intro screen features Mr. Larry's welcome, a mindset-setting exercise "
    "(define Life, Dreams, and Wealth in your own words), and a LinkedIn QR code connecting to Mr. Larry's "
    "professional profile."
)

heading("Parent's Portal — The Command Center", 2)
table(
    ["Feature", "Description", "Game Mechanic"],
    [
        ["Weekly Task Manager", "Create and manage household financial tasks and challenges",
         "Quest board — parent posts missions, child accepts them"],
        ["Weekly Task Allocator", "Assign age-appropriate chores with dollar values to each child",
         "Job board — drag tasks onto child's queue, set pay rates"],
        ["Payroll / Allowance", "Process weekly 'payroll' based on completed tasks",
         "Payday event — animated paycheck generation, teaches employer/employee dynamic"],
        ["Progress Dashboard", "View child's learning progress, simulation scores, and financial habits",
         "Stats screen — XP bars, badges earned, savings growth charts, topic completion map"],
        ["Family Goals", "Set shared family financial goals with visual tracking",
         "Guild quest — shared progress bar, family milestone celebrations"],
        ["Resource Library", "Curated articles, videos, and tools for adult financial literacy",
         "Knowledge base — unlocks as child progresses, so parent learns alongside"],
    ],
    widths=[1.3, 2.2, 3.0],
)

doc.add_paragraph()

heading("Child's Portal — The Game World", 2)
table(
    ["Feature", "Description", "Game Mechanic"],
    [
        ["3D Hallway / World Map", "Navigate the immersive financial world to access topics",
         "Open-world navigation — walk through a 3D hallway or fly over a world map, doors/portals to each topic"],
        ["Weekly Task Manager", "View assigned tasks, mark complete, track earnings",
         "Quest log — accept missions from parent's job board, complete for currency"],
        ["Paystubs / Allowance", "View itemized earnings breakdowns",
         "Inventory — paycheck stub shows gross, deductions, net; teaches pay anatomy through UI"],
        ["Savings Account", "Virtual savings with simulated interest growth",
         "Vault — coins stack up visually, 'interest fairy' adds bonus coins monthly, growth graph unlocks at higher levels"],
        ["Gaming / Shopping", "Simulated marketplace for spending virtual currency",
         "In-game shop — wants vs. needs visualized, impulse buy warnings, budget tracker overlay"],
        ["Investments", "Age-appropriate investment portfolio simulator",
         "Stock market mini-game — buy virtual shares in real companies, track gains/losses, learn diversification"],
    ],
    widths=[1.3, 2.2, 3.0],
)

doc.add_page_break()

# =====================================================================
# 8. GAME DESIGN PHILOSOPHY
# =====================================================================
heading("8. Game Design Philosophy & Mechanics", 1)

body(
    "Game engines (Unity, Unreal, Godot) are more democratized than ever. With tools like Unity's free "
    "tier, Godot's open-source engine, and web-based engines like Three.js and PlayCanvas, we can build "
    "production-quality interactive experiences that rival commercial games — at a fraction of historical "
    "cost. LDW leverages this to create something that feels like a GAME, not a quiz with a game skin."
)

heading("Core Design Principles", 3)

bullet(" Players are dropped into a life scenario and must figure out the financial concepts "
       "through gameplay. The 'lesson' is the experience, not a pre-game lecture.", bp="Learn by Doing: ")
bullet(" Every decision has a visible, traceable consequence. Choose the expensive apartment? "
       "Watch your food budget shrink in real-time. Skip insurance? See what happens when the tree falls.", bp="Consequences, Not Corrections: ")
bullet(" No wrong answers — only different outcomes. The birthday party with cheap cupcakes "
       "is just as fun as the one with the big cake. The game celebrates smart choices, it doesn't punish creative ones.", bp="No Failure States: ")
bullet(" Players FEEL the reward of patience (savings growing) and the sting of impulse "
       "(buyer's remorse animation). Financial behavior is shaped through experience, not rules.", bp="Emotional Resonance: ")
bullet(" Every simulation can be replayed to explore different paths. 'What if I had saved instead "
       "of spent?' The game encourages experimentation without real-world consequences.", bp="Replayability: ")
bullet(" Parents and children play connected games. A parent's payroll decision shows up in the "
       "child's paystub. A child's savings achievement unlocks a family celebration event.", bp="Family Co-Play: ")

doc.add_paragraph()

heading("Game Mechanics Library", 2)
body(
    "LDW uses a diverse set of game mechanics to keep engagement high across ages and learning styles:"
)

doc.add_paragraph()
table(
    ["Mechanic", "Description", "Used In", "Standards Served"],
    [
        ["Branching Narrative", "Story paths diverge based on player decisions. Multiple endings per scenario. Each branch teaches different concepts.",
         "Every simulation", "FDM (core mechanic)"],
        ["Resource Management", "Players manage limited currency/budget across competing needs. Visual resource bars drain and fill in real-time.",
         "Birthday Budget, Grocery Store, Apartment Hunt, Life After Graduation", "SS, FDM"],
        ["Timed Events", "Random life events interrupt gameplay (car breaks down, sale opportunity, weather). Players must react under time pressure.",
         "Lemonade Stand, Class Trip Fund, Apartment Hunt, Insurance Gamble", "RM, FDM"],
        ["Growth Visualizer", "Animated compound growth — money 'grows' like a plant, interest 'rains' down. Time-lapse slider shows 1yr, 10yr, 30yr.",
         "Savings Snowball, Investment Challenge, Generational Wealth", "INV, SS"],
        ["Dual Comparison", "Side-by-side split screen showing two paths simultaneously (insured vs. uninsured, jar vs. bank, college vs. trade).",
         "Savings Snowball, Insurance Gamble, College vs. Career, Community Resource Map", "All domains"],
        ["Quest Lines", "Goal-tracking with milestones, progress bars, and celebration events. Short/intermediate/long quests mirror real goal horizons.",
         "Three Piggy Banks, Class Trip Fund, all DREAMS-phase sims", "SS, FDM"],
        ["Economy Simulation", "Player-run businesses with revenue, costs, profit, competition, and scaling decisions.",
         "Lemonade Stand, Neighborhood Business, Shark Tank Jr.", "EI, FDM"],
        ["NPC Mentorship", "Mr. Larry appears as a guide character throughout — video intros, in-game tips, celebratory moments, and 'ask Mr. Larry' help buttons.",
         "Every simulation", "All domains"],
        ["Social Reality Meter", "Shows the gap between social media appearance and financial reality. Players post a 'social media update' and see the contrast.",
         "Apartment Hunt, Life After Graduation", "FDM, SS"],
        ["Map Exploration", "Navigate neighborhood/community maps to discover resources (or discover they're missing). Click on buildings to learn about services.",
         "Community Resource Map, Generational Wealth Gap", "FDM, CD, SS"],
        ["Portfolio Builder", "Drag-and-drop investment allocation across asset classes. Real-time visualization of risk/return tradeoffs.",
         "Investment Challenge, Generational Wealth", "INV"],
        ["Tax Calculator", "Step-by-step guided tax form with real 1040 fields. Deductions and credits unlock like power-ups.",
         "The Tax Return", "EI, FDM"],
        ["Legacy Tracker", "Multi-generational wealth visualization showing how decisions cascade across generations. Family tree with wealth overlay.",
         "Generational Wealth Gap, Building Generational Wealth", "INV, FDM"],
        ["Badge & Achievement System", "Collectible badges for completing simulations, mastering concepts, maintaining savings streaks, and family co-play milestones.",
         "All simulations", "Motivation & engagement"],
    ],
    widths=[1.1, 2.2, 1.5, 1.0],
)

doc.add_page_break()

# =====================================================================
# 9. THE 3D WORLD
# =====================================================================
heading("9. The 3D World: Navigating Financial Life", 1)

body(
    "The primary navigation experience is an immersive 3D environment — Mr. Larry's original vision "
    "of a hallway with numbered subjects on the walls, now brought to life as a full game world."
)

heading("The Financial Hallway", 3)
body(
    "Players walk through a 3D corridor. The walls are richly textured — warm wood panels and clean white "
    "subject boards, exactly as Mr. Larry sketched. Each board displays a subject number (#1 through #13), "
    "a title, and a status indicator (locked, available, in-progress, completed). As players complete "
    "topics, the hallway transforms — lights brighten, the iceberg progress meter rises, and completed "
    "boards show a golden frame."
)

heading("World Variations by Grade Band", 3)
table(
    ["Grade Band", "World Theme", "Navigation Style", "Visual Tone"],
    [
        ["K-2", "The Treehouse Village", "Walk between colorful treehouses, each one a lesson. Friendly animal guides.",
         "Bright, warm, rounded shapes. Think Animal Crossing meets Sesame Street."],
        ["3-5", "The Town Square", "A small town with a grocery store, bank, school, park. Each building is a simulation.",
         "Clean, modern, inviting. Real-world buildings that feel familiar and safe."],
        ["6-8", "The City Block", "An urban neighborhood with apartments, businesses, banks, and community centers. More realistic.",
         "Stylized realism. City energy. Day/night cycle. Weather events affect gameplay."],
        ["9-12", "The Life Map", "A zoomed-out world map showing career paths, cities, universities, and life milestones as destinations.",
         "Clean data visualization meets open world. Professional, aspirational, future-oriented."],
    ],
    widths=[0.8, 1.3, 2.5, 2.0],
)

doc.add_paragraph()

heading("The Iceberg Progress System", 3)
body(
    "A persistent iceberg sits in the corner of the screen (or at the center of the world hub). "
    "At the start, only the tip shows above water. As players complete simulations and demonstrate "
    "financial knowledge, more of the iceberg rises — revealing layers labeled with financial concepts: "
    "Budgeting, Saving, Credit, Investing, Insurance, Generational Wealth. When the full iceberg is "
    "visible, the player has 'surfaced' their complete financial foundation. This is the ultimate "
    "progress tracker — and a visual metaphor Mr. Larry has used from the beginning."
)

doc.add_page_break()

# =====================================================================
# 10. SIMULATION LESSONS
# =====================================================================
heading("10. The Game Engine: Phenomenon-Based Simulation Lessons", 1)

body(
    "Each simulation follows Mr. Larry's core principle: start with a life experience, let the player "
    "discover the financial concept through gameplay, then name it. This is phenomenon-based learning — "
    "the same methodology used in Next Generation Science Standards (NGSS), now applied to financial literacy."
)

heading("Simulation Structure", 3)
body("Every simulation follows a consistent 4-week gameplay rhythm:")

doc.add_paragraph()
table(
    ["Week", "Phase", "Gameplay", "Learning Mode"],
    [
        ["Week 1", "DISCOVER", "Mr. Larry intro video. Player enters the scenario world. Explores the environment. "
         "Makes predictions. Pre-assessment disguised as a 'character questionnaire.'",
         "Anchoring phenomenon presented. Driving question posed. Prior knowledge activated."],
        ["Week 2", "PLAY", "Full branching simulation. Multiple decision points with visible consequences. "
         "Replayable for different paths. Decision journal auto-tracks choices and outcomes.",
         "Core concepts discovered through gameplay. Standards addressed through authentic problem-solving."],
        ["Week 3", "FAMILY QUEST", "Shared parent-child challenge. Real-world application of the simulation's concepts. "
         "Conversation starters. Photo/journal uploads to the family portfolio.",
         "Transfer from virtual to real world. Family as learning unit. Social-emotional connection to content."],
        ["Week 4", "LEVEL UP", "Reflection quest. Assessment disguised as a 'knowledge challenge.' Goal-setting for "
         "next simulation. Badge awarded. Iceberg rises. New world area unlocked.",
         "Metacognition. Self-assessment. Goal-setting skill practice. Celebration and motivation."],
    ],
    widths=[0.6, 0.8, 2.7, 2.4],
)

doc.add_page_break()

# =====================================================================
# 11. COMPLETE SIMULATION LIBRARY
# =====================================================================
heading("11. Complete Simulation Library (K-12)", 1)

body(
    "20 phenomenon-based simulations, each anchored in a real-world financial life experience. "
    "Every simulation is a complete 4-week gameplay cycle with branching decisions, family quests, "
    "and standards-aligned assessment."
)

# K-2
heading("K-2: The Treehouse Village (4 Simulations)", 2)

sims_k2 = [
    ["K2-01", "The Birthday Party Budget", "LIFE",
     "Maya has $20 for her party but wants everything — cake, balloons, games, AND favors. What should she choose?",
     "How do we decide what to buy when we can't buy everything?",
     "Wants vs. needs, opportunity cost, resource allocation, prioritization",
     "Resource management with animated money jar. 3 party outcomes — all fun. Teaches that constraint breeds creativity."],
    ["K2-02", "The Lemonade Stand", "DREAMS",
     "Jayden wants a $15 soccer ball. Mom says earn it. But he needs $5 for supplies first. You need money to MAKE money?",
     "Why do you need to spend money before you can earn money?",
     "Revenue, costs, profit, entrepreneurship basics, weather as market factor",
     "Business simulation with customer counter, weather events, location choices. Visual P&L at checkout."],
    ["K2-03", "The Coin Quest", "LIFE",
     "Grandma gives a jar of mixed coins. The big ones aren't always worth more. A dime beats a nickel. How much is there?",
     "Why isn't a bigger coin always worth more money?",
     "Coin identification, counting, making change, multiple pathways to same total",
     "Drag-and-drop coin sorting. Three store visits requiring exact change. Final purchase decision."],
    ["K2-04", "The Three Piggy Banks", "DREAMS",
     "$3/week split into Save, Spend, Share. After 4 weeks, who reaches their dream toy depends entirely on how they split.",
     "What happens to your money when you split it into Save, Spend, and Share?",
     "Allocation, saving toward goals, delayed gratification, giving",
     "Animated piggy banks with 4-week cycle. Temptation events. Windfall decisions. Visual savings growth."],
]

for sim in sims_k2:
    heading(f"{sim[0]}: {sim[1]} [{sim[2]}]", 3)
    bullet(sim[3], bp="Phenomenon: ")
    bullet(sim[4], bp="Driving Question: ")
    bullet(sim[5], bp="Concepts: ")
    bullet(sim[6], bp="Game Design: ")
    doc.add_paragraph()

doc.add_page_break()

# 3-5
heading("3-5: The Town Square (6 Simulations)", 2)

sims_35 = [
    ["35-01", "HealthMart Foods: The Grocery Challenge", "LIFE",
     "Mom sends you with $40 and a list. Every item has 3 versions: cheap, moderate, premium. Price, health, and prep time compete.",
     "How do you make smart choices when every option has a different price, quality, and trade-off?",
     "Comparison shopping, unit pricing, budgeting, nutrition economics, impulse control",
     "FLAGSHIP SIM. Virtual grocery store with live budget/health/time trackers. Coupon hunting. Impulse sale events. Checkout summary with Mr. Larry commentary."],
    ["35-02", "The Class Trip Fund", "DREAMS",
     "Class trip costs $375. School pays half. You have 8 weeks to raise $187.50. Bake sale? Car wash? Read-a-thon? Rain ruins your plan.",
     "How do groups work together to reach a big financial goal?",
     "Collaborative goal-setting, fundraising economics, contingency planning, revenue vs. cost",
     "8-week fundraising campaign sim. Thermometer progress bar. Weather events. Windfall management. Multiple strategies to compare."],
    ["35-03", "Paycheck Payday", "LIFE",
     "Your chore paycheck is $25. After 'community tax' ($2.50) and 'family fund' ($1.00), you take home $21.50. Why less than you earned?",
     "Why is the money you take home different from the money you earned?",
     "Gross vs. net income, deductions, taxes, fixed expenses, budgeting from net",
     "Animated paycheck generator. Deduction reveal with Mr. Larry explanation. Budget allocator. Streaming 'bill' teaches fixed expenses."],
    ["35-04", "The Savings Snowball", "DREAMS/WEALTH",
     "Twins save $5/week. Amara uses a jar. Zuri uses a savings account with 2% interest. 6 months later, Zuri has extra money she never earned.",
     "How does money grow when you're not even touching it?",
     "Interest, compound growth, savings strategies, time value of money",
     "Side-by-side savings race. 'Interest fairy' animation. Time-lapse slider: 1yr, 5yr, 10yr. Dramatic gap visualization at scale."],
    ["35-05", "The Neighborhood Business", "DREAMS",
     "Three kids start businesses: dog walking ($120 revenue), lawn care ($80 revenue, $30 expenses), tutoring ($60, no expenses). Who profited most?",
     "Why is earning the most money not the same as keeping the most money?",
     "Revenue vs. profit, business expenses, pricing strategy, competition",
     "Choose and manage a business for 4 weeks. P&L dashboard builds weekly. Competition events. Scaling decisions."],
    ["35-06", "The Community Resource Map", "WEALTH",
     "Two families, same income. Family A lives near a credit union and community garden. Family B lives near check-cashing stores and payday lenders. After 1 year: $2,400 saved vs. $200. Why?",
     "How does where you live affect how much money you can save?",
     "Financial access, predatory vs. helpful services, community resources, systemic barriers",
     "Dual-neighborhood map navigation. Manage both families' finances. Watch the gap grow month by month. Identify 'hidden costs' of under-resourced communities."],
]

for sim in sims_35:
    heading(f"{sim[0]}: {sim[1]} [{sim[2]}]", 3)
    bullet(sim[3], bp="Phenomenon: ")
    bullet(sim[4], bp="Driving Question: ")
    bullet(sim[5], bp="Concepts: ")
    bullet(sim[6], bp="Game Design: ")
    doc.add_paragraph()

doc.add_page_break()

# 6-8
heading("6-8: The City Block (6 Simulations)", 2)

sims_68 = [
    ["68-01", "The Credit Score Mystery", "LIFE",
     "Destiny gets a car loan at 4%. Marcus gets offered 18%. Same bank. Both work. The difference? A three-digit number neither understands.",
     "What is a credit score, and why does it control so much of your financial life?",
     "Credit scores (5 factors), borrowing costs, credit building vs. damaging, long-term cost of rates",
     "2-year dual-character management. Live credit score meter (300-850). Every decision moves the needle. Final loan comparison: same car, up to $4,000 difference."],
    ["68-02", "The Investment Challenge", "WEALTH",
     "In 1990, Alex puts $1,000 in savings (2%/yr). Jordan puts $1,000 in an index fund (10%/yr avg). In 2026: $2,039 vs. $30,912. How?",
     "How do small differences in where you put money create huge differences over time?",
     "Compound interest, stocks/bonds/funds, risk vs. return, market volatility, panic selling",
     "30-year investment sim with real historical events (2008 crash, 2020 dip, AI boom). Panic-sell or hold steady. Time-lapse portfolio growth. The most dramatic visualization in the game."],
    ["68-03", "The Apartment Hunt", "LIFE",
     "Three friends earn $2,500/month. Apartments: $1,500 (looks amazing on Instagram), $750 (decent), $500 (basic). The Instagram friend can't afford groceries.",
     "How much should go to housing, and what happens when you get that ratio wrong?",
     "50/30/20 rule, fixed vs. variable expenses, budgeting, emergency funds, social media vs. reality",
     "Full monthly budget management. Random life events. Social media reality meter — post your 'highlight reel' and see the truth underneath."],
    ["68-04", "Shark Tank Jr.: The Pitch", "DREAMS",
     "A 13-year-old made $450 selling fidget phone cases. She needs $2,000 to scale. Loan? Investor? Crowdfund? Bootstrap?",
     "How do entrepreneurs fund growth, and what does each funding source really cost?",
     "Business funding, equity, interest, break-even, profit margins, pitch skills",
     "Product development → funding decision → 6-month operations → Shark Tank pitch. Four funding paths with real financial consequences. Pitch recorder for presentations."],
    ["68-05", "The Insurance Gamble", "WEALTH",
     "Family A pays $400/month insurance. Family B pays $0, saving $14,400 over 3 years. Then a tree falls on both houses. Family A pays $1,000. Family B owes $45,000.",
     "Is insurance a waste of money — or the smartest thing you'll ever pay for?",
     "Insurance purpose and types, premiums vs. deductibles, risk probability, cost-benefit analysis",
     "5-year dual-household sim. Random event generator each year. Probability visualizer shows: it's not IF, it's WHEN. Coverage level adjuster."],
    ["68-06", "The Generational Wealth Gap", "WEALTH",
     "1960: Johnson family gets a VA home loan. Williams family is denied (redlining). Same income. By 2026: $450K wealth gap. Same work ethic. Different rules.",
     "How did historical policies create today's wealth gap, and what can be done?",
     "Generational wealth, redlining, systemic barriers, compound effect of homeownership, community solutions",
     "Three-generation family simulator (1960→1990→2026). Policy reveals change the options available. Solutions explorer at end. Historical context without blame — systems focus."],
]

for sim in sims_68:
    heading(f"{sim[0]}: {sim[1]} [{sim[2]}]", 3)
    bullet(sim[3], bp="Phenomenon: ")
    bullet(sim[4], bp="Driving Question: ")
    bullet(sim[5], bp="Concepts: ")
    bullet(sim[6], bp="Game Design: ")
    doc.add_paragraph()

doc.add_page_break()

# 9-12
heading("9-12: The Life Map (4 Simulations)", 2)

sims_912 = [
    ["912-01", "The Life After Graduation Simulator", "LIFE",
     "June. You graduated. $35K job offer. $12K student loans. $1,500 in the bank. Apartment deposit due August 1. First paycheck: August 15. Welcome to adulting.",
     "What does it actually cost to start your adult life?",
     "Post-grad financial planning, take-home pay, housing ratio, loan repayment, emergency fund, real budgeting",
     "THE most realistic sim. Day-by-day 90-day survival. Real checking account interface. Overdraft warnings. Adulting surprises (parking ticket, work lunch, friend's wedding gift). Three housing paths cascade through everything."],
    ["912-02", "The Tax Return", "LIFE",
     "You earned $28,000. Employer took $3,360 in taxes. Your coworker with a kid gets $2,100 back. You get $400. Same job, same pay, different life circumstances.",
     "How do taxes actually work, and how can understanding them put money in your pocket?",
     "Tax brackets, W-2, deductions vs. credits, EITC, student loan interest, filing status effects",
     "Step-by-step 1040 tax return with real fields. Three-scenario side-by-side: single, single parent, married. Credits unlock like power-ups. Refund strategy decision at end."],
    ["912-03", "College vs. Career: The $100K Decision", "DREAMS",
     "Twins graduate. Aaliyah goes to university ($25K/yr). Devon starts an apprenticeship ($35K immediately). At 26, Aaliyah has a degree + $60K debt. Devon has 4 years of experience + $0 debt. At 65, who has more wealth?",
     "Is college always worth the investment?",
     "Education ROI, opportunity cost, student loans, career earning curves, lifetime wealth modeling",
     "40-year life financial modeler. Six career paths: trades, military, community college, state university, private university, entrepreneurship. Real salary data. The answer depends on the path — no universal 'right choice.'"],
    ["912-04", "Building Generational Wealth", "WEALTH",
     "Great-grandmother saves $50/month starting in 1970. Index fund at 10%/yr. Never touches it. By 2026: over $1.2 million. One decision, three generations, millionaire family. What if she started with $25? What if she started 10 years later?",
     "How do ordinary families build extraordinary wealth across generations?",
     "Compound growth at scale, diversification, estate planning basics, legacy vision, values-driven finance",
     "CAPSTONE SIM. Three-generation wealth management. Set savings, choose investments, navigate life events, decide how wealth transfers. Ends with a personal Generational Wealth Blueprint and Family Legacy Letter."],
]

for sim in sims_912:
    heading(f"{sim[0]}: {sim[1]} [{sim[2]}]", 3)
    bullet(sim[3], bp="Phenomenon: ")
    bullet(sim[4], bp="Driving Question: ")
    bullet(sim[5], bp="Concepts: ")
    bullet(sim[6], bp="Game Design: ")
    doc.add_paragraph()

doc.add_page_break()

# =====================================================================
# 12. FAMILY GAMEPLAY
# =====================================================================
heading("12. Family Gameplay: The Parent-Child Connection", 1)

body(
    "Every simulation includes a Week 3 'Family Quest' — a real-world challenge that parents and children "
    "complete together. This is where the game bridges virtual learning and actual life."
)

doc.add_paragraph()
table(
    ["Simulation", "Family Quest", "Real-World Connection"],
    [
        ["Birthday Party Budget", "Plan a real family movie night on a $10 budget",
         "Child practices budgeting with real money or play money. Parent guides without deciding."],
        ["The Lemonade Stand", "Set up a real mini-business for a day (lemonade, bake sale, car wash)",
         "Child tracks real costs, revenue, profit. Deposits 'earnings' into savings jar."],
        ["The Coin Quest", "Collect loose change around the house for a week, sort and count together",
         "Decision: spend it, save it, or share it? Photo of sorted coins."],
        ["Three Piggy Banks", "Set up 3 real jars (Save, Spend, Share). Split allowance weekly for 4 weeks.",
         "Parent photographs jars weekly. Celebrate the SAVE achievement."],
        ["HealthMart Foods", "Real grocery trip with a $40 budget. Child makes the list, comparison shops, tracks spending.",
         "Family cooks a meal from purchased items. Discussion: What surprised you about prices?"],
        ["Class Trip Fund", "Identify a family financial goal. Create a thermometer chart. Save toward it for 4 weeks.",
         "Child proposes ways to earn extra toward the goal."],
        ["Paycheck Payday", "Parent shares a simplified household bill. Child budgets allowance including a 'bill' payment.",
         "Discussion: What are our family's biggest expenses?"],
        ["Savings Snowball", "Open a real savings account (or virtual). Contribute weekly. Track interest.",
         "Discussion: Where else does compound growth happen in life?"],
        ["Credit Score Mystery", "Research one action that builds credit and one that damages it.",
         "Older teens: pull free credit report together at annualcreditreport.com."],
        ["Investment Challenge", "Pick 3 companies the family uses daily. Look up stock prices. Would $100 last year have grown?",
         "Discussion: Why does a company's stock price change?"],
        ["Apartment Hunt", "Parent shares approximate monthly budget categories. Calculate housing percentage.",
         "Discussion: Does our family follow the 50/30/20 rule?"],
        ["Insurance Gamble", "Identify all family insurance. Calculate total monthly cost.",
         "Discussion: What would happen if we didn't have health insurance and someone got sick?"],
        ["Life After Graduation", "Calculate the real cost of first 90 days after graduation. Include deposit, rent, wardrobe, food.",
         "Discussion: What did we wish we'd known when we started out?"],
        ["Building Generational Wealth", "Write a Family Legacy Letter about financial values, lessons, and hopes for the next generation.",
         "Start a real family investment plan, even $10/month."],
    ],
    widths=[1.3, 2.5, 2.7],
)

doc.add_page_break()

# =====================================================================
# 13. ASSESSMENT
# =====================================================================
heading("13. Assessment Through Play", 1)

body(
    "LDW does not use traditional tests. Assessment is embedded in gameplay — every decision a player "
    "makes IS the assessment. The system tracks decision patterns, financial outcomes, and growth over time."
)

heading("Assessment Layers", 3)
table(
    ["Layer", "What It Measures", "How It Works", "Visible To"],
    [
        ["Decision Analytics", "Quality of financial decisions across simulations",
         "Every branching choice is logged and scored against optimal financial outcomes. Patterns reveal strengths and growth areas.",
         "Parent Dashboard, Teacher Dashboard (if school-deployed)"],
        ["Simulation Outcomes", "Whether players achieve financial goals within scenarios",
         "Did the player stay within budget? Reach the savings goal? Build good credit? Each sim has measurable success criteria.",
         "Player (as game score), Parent, Teacher"],
        ["Replay Improvement", "Growth over multiple playthroughs",
         "Players who replay and make better choices demonstrate learning transfer. The system tracks improvement trajectories.",
         "Player (progress bar), Parent"],
        ["Family Quest Completion", "Real-world application of concepts",
         "Photo/journal uploads from family quests. Self-reported completion. Parent confirmation.",
         "Family portfolio, Parent"],
        ["Knowledge Challenges", "Concept mastery verification",
         "End-of-simulation quizzes disguised as 'Level Up Challenges.' Gamified assessment with point rewards.",
         "Player, Parent, Teacher"],
        ["Portfolio Review", "Cumulative financial literacy growth",
         "Semester/year-end portfolio of all simulation results, family quests, and growth trajectory. Printable for schools.",
         "All stakeholders"],
    ],
    widths=[1.1, 1.5, 2.2, 1.7],
)

doc.add_paragraph()

heading("Standards Reporting", 3)
body(
    "For school deployments, LDW generates standards-aligned reports mapping each student's performance "
    "to Jump$tart and CEE standards. Teachers can see at a glance which standards each student has "
    "demonstrated proficiency in, and which need reinforcement. This makes LDW not just a game, "
    "but a standards-compliant assessment tool."
)

doc.add_page_break()

# =====================================================================
# 14. PRIOR WORK
# =====================================================================
heading("14. Prior Work & Existing Assets", 1)

body(
    "Significant development has already been completed, providing a strong foundation:"
)

doc.add_paragraph()
table(
    ["Asset", "Status", "Description", "Integration Plan"],
    [
        ["HealthMart Foods Nutrition Simulation", "Complete",
         "Adobe Captivate branching simulation — grocery shopping with price/health/time decisions. Full SCORM package.",
         "Evolve into the flagship 35-01 Grocery Challenge. Expand branches, add game mechanics, modernize UI."],
        ["VR / 360 Simulation", "Complete",
         "SCORM 1.2 Captivate sim with panoramic environment and quiz interactions.",
         "Reference for 3D hallway design. Technology approach informs immersive world-building."],
        ["Simulation Branching Logic", "Complete",
         "Detailed Excel mapping all decision branches, scoring rubrics, and outcome paths.",
         "Blueprint for expanding branching logic to all 20 simulations."],
        ["Parent & Child Portal Mockups", "Prototype",
         "Functional mockups: Parent Portal (Task Manager, Allocator, Payroll) and Child Portal (Tasks, Savings, Shopping, Investments).",
         "Direct input for UI/UX game design. Feature list validated by Mr. Larry."],
        ["360 Interactive eLearning Prototype", "Complete",
         "16-interaction-type prototype with draggable panorama, assessments, calculators, flashcards, gamification.",
         "Demonstrates proof of concept for interactive game mechanics. Reference for interaction design patterns."],
        ["Logo & Brand Assets", "Complete",
         "19 logo variations, character avatars, iceberg concept graphics, social media templates.",
         "Brand identity for the game. Iceberg becomes the core progress visualization."],
        ["Presentation Decks & Pitch Materials", "Complete",
         "PowerPoint presentations covering concept, competitive positioning, and business model.",
         "Context for game narrative and Mr. Larry's vision. Competitive differentiators inform game design."],
        ["Meeting Recordings & Transcripts", "Complete",
         "July 2022 meetings with Amanda (NDAs signed). March 2026 planning session with new vision.",
         "Historical context and design requirements. Updated vision informs all current development."],
    ],
    widths=[1.3, 0.7, 2.3, 2.2],
)

doc.add_page_break()

# =====================================================================
# 15. NEXT STEPS
# =====================================================================
heading("15. Next Steps", 1)

body("We recommend the following actions to move from proposal to product:")

doc.add_paragraph()

steps = [
    ("Review & Align: ", "Mr. Wimsatt reviews this proposal. Are the 20 simulations the right life experiences? "
     "Are there scenarios missing from his vision? Does the game design match what he sees?"),
    ("Standards Deep Dive: ", "Select target states for initial launch (California's 2027 mandate is the nearest opportunity). "
     "Complete formal standards crosswalk with specific state benchmarks for adoption approval."),
    ("Curriculum Detailing: ", "Deep-dive into the first 4 simulations (K2-01 through K2-04). Write full branching scripts, "
     "dialogue, decision trees, scoring rubrics, and family quest guides."),
    ("Game Design Document: ", "Create a formal GDD (Game Design Document) for the chosen game engine. Define art style, "
     "character design (Mr. Larry's avatar), UI/UX wireframes, and interaction patterns."),
    ("Prototype Build: ", "Build a playable prototype of one simulation (recommend 35-01: HealthMart Foods, since prior work "
     "exists). Demonstrate the full 4-week gameplay cycle with real branching, family quest, and assessment."),
    ("Pilot Families: ", "Identify 5-10 pilot families to play the prototype. Collect feedback on engagement, "
     "learning outcomes, family co-play experience, and age-appropriateness."),
    ("Advisory Board: ", "Assemble an advisory board: financial literacy educators, K-12 teachers, game designers, "
     "and community leaders from target populations to guide development."),
]

for i, (bp, txt) in enumerate(steps, 1):
    bullet(txt, bp=f"{i}. {bp}")

doc.add_paragraph()
doc.add_paragraph()

# Closing
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("_" * 60)
r.font.color.rgb = TEAL
r.font.size = Pt(10)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Alexandria's Design LLC")
r.bold = True
r.font.size = Pt(14)
r.font.color.rgb = NAVY
r.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Dr. Charles Martin | 310-709-4893 | alexandriasworld1234@gmail.com")
r.font.size = Pt(11)
r.font.color.rgb = MID_GRAY
r.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('"Where Families Build Financial Futures Together"')
r.italic = True
r.font.size = Pt(12)
r.font.color.rgb = DARK_GOLD
r.font.name = 'Calibri'

# Save
doc.save(OUT)
print(f"Saved: {OUT}")
print(f"Size: {os.path.getsize(OUT) / 1024:.0f} KB")
