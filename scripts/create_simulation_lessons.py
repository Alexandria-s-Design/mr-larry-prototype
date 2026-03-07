"""
Generate phenomenon-based financial literacy simulation lessons for Life Dreams Wealth.
Organized by grade band, mapped to Jump$tart / CEE national standards.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

REPO = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(REPO, "LDW_Phenomenon_Based_Simulation_Lessons.docx")

# Brand colors
NAVY = RGBColor(0x0F, 0x0F, 0x1A)
GOLD = RGBColor(0xC9, 0xA8, 0x4C)
DARK_GOLD = RGBColor(0xA0, 0x85, 0x3D)
TEAL = RGBColor(0x2A, 0x9D, 0x8F)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK_TEXT = RGBColor(0x1A, 0x1A, 0x2E)
MID_GRAY = RGBColor(0x55, 0x55, 0x55)
CORAL = RGBColor(0xE0, 0x6C, 0x5A)
PURPLE = RGBColor(0x6B, 0x5B, 0x95)

# Phase colors for LDW framework
LIFE_COLOR = "2A9D8F"    # Teal
DREAMS_COLOR = "C9A84C"  # Gold
WEALTH_COLOR = "6B5B95"  # Purple

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = DARK_TEXT
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        if level == 1:
            run.font.color.rgb = NAVY
            run.font.size = Pt(22)
        elif level == 2:
            run.font.color.rgb = DARK_GOLD
            run.font.size = Pt(16)
        elif level == 3:
            run.font.color.rgb = TEAL
            run.font.size = Pt(13)
        run.font.name = 'Calibri'
    return h


def add_body(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run.font.color.rgb = DARK_TEXT
    return p


def add_bullet(text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        run.font.color.rgb = DARK_TEXT
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        run.font.color.rgb = DARK_TEXT
    return p


def add_table(headers, rows, col_widths=None, header_color="0F0F1A"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            run.font.color.rgb = DARK_TEXT
            if r_idx % 2 == 0:
                set_cell_shading(cell, "F5F5F0")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table


def add_sim_lesson(num, title, grade, phenomenon, driving_question,
                   standards, ldw_phase, sim_description, branches,
                   family_activity, assessment, parent_portal, child_portal,
                   phase_color):
    """Add a full phenomenon-based simulation lesson."""

    # Lesson header with phase color band
    h = doc.add_heading(f"Simulation {num}: {title}", level=2)
    for run in h.runs:
        run.font.name = 'Calibri'

    # Grade & metadata bar
    meta_table = doc.add_table(rows=1, cols=4)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.style = 'Table Grid'
    meta_data = [
        ("Grade Band", grade),
        ("LDW Phase", ldw_phase),
        ("Duration", "4 Weeks"),
        ("Type", "Branching Simulation"),
    ]
    for i, (label, value) in enumerate(meta_data):
        cell = meta_table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(f"{label}: ")
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Calibri'
        run.font.color.rgb = WHITE
        run = p.add_run(value)
        run.font.size = Pt(9)
        run.font.name = 'Calibri'
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, phase_color)

    doc.add_paragraph()

    # Anchoring Phenomenon
    add_heading("Anchoring Phenomenon", 3)
    p = add_body(phenomenon)
    p.runs[0].italic = True

    # Driving Question
    add_heading("Driving Question", 3)
    p = add_body(f'"{driving_question}"')
    p.runs[0].bold = True

    # Standards Alignment
    add_heading("Financial Literacy Standards", 3)
    for std in standards:
        add_bullet(std)

    # Simulation Description
    add_heading("Simulation Overview", 3)
    add_body(sim_description)

    # Branching Paths
    add_heading("Decision Branches", 3)
    add_body("Students encounter the following decision points during the simulation:")
    doc.add_paragraph()
    branch_headers = ["Decision Point", "Option A", "Option B", "Option C (if applicable)", "Consequence"]
    add_table(branch_headers, branches, col_widths=[1.2, 1.2, 1.2, 1.2, 1.7])
    doc.add_paragraph()

    # 4-Week Rhythm
    add_heading("4-Week Lesson Rhythm", 3)
    week_data = [
        ["Week 1: DISCOVER", f"Mr. Larry introduces the phenomenon through a short video. Students explore the scenario, make predictions, and take a pre-assessment. The anchoring question is posed: \"{driving_question}\""],
        ["Week 2: SIMULATE", f"Students run the full branching simulation. Each decision point reveals consequences. Students track their choices and outcomes on a decision journal. Multiple playthroughs encouraged to explore different paths."],
        ["Week 3: FAMILY CONNECT", family_activity],
        ["Week 4: REFLECT & APPLY", assessment],
    ]
    add_table(["Week", "Activities"], week_data, col_widths=[1.3, 5.2])
    doc.add_paragraph()

    # Portal Integration
    add_heading("Portal Integration", 3)
    add_bullet(parent_portal, bold_prefix="Parent Portal: ")
    add_bullet(child_portal, bold_prefix="Child Portal: ")

    doc.add_page_break()


# ============================================================
# COVER PAGE
# ============================================================

for _ in range(4):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("LIFE DREAMS WEALTH")
run.font.size = Pt(36)
run.font.color.rgb = NAVY
run.bold = True
run.font.name = 'Calibri'

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("Phenomenon-Based Simulation Lessons")
run.font.size = Pt(20)
run.font.color.rgb = DARK_GOLD
run.font.name = 'Calibri'

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub2.add_run("Standards-Aligned Financial Literacy Curriculum")
run.font.size = Pt(14)
run.font.color.rgb = TEAL
run.font.name = 'Calibri'

doc.add_paragraph()

divider = doc.add_paragraph()
divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = divider.add_run("_" * 60)
run.font.color.rgb = TEAL
run.font.size = Pt(10)

doc.add_paragraph()

details = [
    "Prepared for Mr. Lawrence Wimsatt | Life Dreams Wealth",
    "Prepared by Alexandria's Design LLC | Dr. Charles Martin",
    "March 2026 | Version 1.0",
    "",
    "24 Phenomenon-Based Simulations Across K-12",
    "Mapped to Jump$tart & CEE National Standards",
]
for d in details:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(d)
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_TEXT if d else WHITE
    run.font.name = 'Calibri'

doc.add_page_break()

# ============================================================
# STANDARDS FRAMEWORK
# ============================================================

add_heading("Financial Literacy Standards Framework", 1)

add_body(
    "All simulation lessons are mapped to the nationally recognized standards "
    "from the Jump$tart Coalition for Personal Financial Literacy and the "
    "Council for Economic Education (CEE). These six domains form the backbone "
    "of the LDW curriculum:"
)

doc.add_paragraph()

std_headers = ["Domain", "Code", "Description", "LDW Phase"]
std_rows = [
    ["Spending & Saving", "SS", "Budgeting, comparison shopping, opportunity cost, saving strategies", "LIFE"],
    ["Credit & Debt", "CD", "Understanding credit, responsible borrowing, debt management, credit scores", "LIFE / WEALTH"],
    ["Employment & Income", "EI", "Earning income, career planning, taxes, benefits, pay structures", "LIFE"],
    ["Investing", "INV", "Risk/return, compound growth, diversification, investment vehicles", "WEALTH"],
    ["Risk Management & Insurance", "RM", "Protecting assets, insurance types, emergency preparedness", "DREAMS / WEALTH"],
    ["Financial Decision Making", "FDM", "Goal setting, planning, evaluating tradeoffs, behavioral finance", "DREAMS"],
]
add_table(std_headers, std_rows, col_widths=[1.5, 0.6, 2.8, 1.6])

doc.add_paragraph()

add_body(
    "Each simulation targets 2-3 standards domains. The phenomenon-based approach means students "
    "encounter a real-world financial scenario first, then discover the underlying concepts through "
    "exploration, decision-making, and consequence — just like Mr. Larry's original grocery store "
    "simulation, but expanded across all financial literacy domains."
)

doc.add_page_break()

# ============================================================
# GRADE BAND OVERVIEW
# ============================================================

add_heading("Grade Band Overview", 1)

add_body("Simulations are designed in four developmental tiers:")

doc.add_paragraph()

band_headers = ["Grade Band", "# Sims", "Developmental Focus", "Simulation Style"]
band_rows = [
    ["K-2 (Ages 5-8)", "4", "Concrete, visual, narrative-driven. Identify coins/bills, wants vs. needs, sharing & saving.",
     "Guided story simulations with simple 2-way choices. Animated characters. Audio narration."],
    ["3-5 (Ages 8-11)", "6", "Applied math, budgeting basics, earning & spending, goal-setting, community economics.",
     "Interactive branching scenarios with 3-way choices. Calculator tools. Visual dashboards."],
    ["6-8 (Ages 11-14)", "8", "Abstract financial concepts, percentages, compound interest, credit, entrepreneurship, systemic inequity.",
     "Complex multi-path simulations. Spreadsheet-style trackers. Real-world data integration. Multiplayer elements."],
    ["9-12 (Ages 14-18)", "6", "Advanced investing, taxes, insurance, college/career financial planning, generational wealth.",
     "Full financial modeling simulations. Stock market games. Life scenario engines. Portfolio management."],
]
add_table(band_headers, band_rows, col_widths=[1.2, 0.5, 2.5, 2.3])

doc.add_paragraph()

add_heading("Alignment to LDW Framework", 3)
add_body(
    "Every grade band covers all three LDW phases (Life, Dreams, Wealth) at age-appropriate depth. "
    "A kindergartner exploring \"wants vs. needs\" is doing LIFE-phase work. A 5th grader setting a "
    "savings goal for a class trip is doing DREAMS-phase work. An 8th grader running an investment "
    "simulator is doing WEALTH-phase work. The framework scales."
)

doc.add_page_break()

# ============================================================
# K-2 SIMULATIONS
# ============================================================

add_heading("Grade Band: K-2 (Ages 5-8)", 1)
add_body(
    "Young learners experience financial concepts through story-driven simulations with familiar characters, "
    "colorful visuals, and simple two-choice decisions. Mr. Larry narrates each scenario in a warm, "
    "encouraging voice. Parents receive companion activities for home."
)

doc.add_paragraph()

# K-2 Sim 1
add_sim_lesson(
    num="K2-01",
    title="The Birthday Party Budget",
    grade="K-2",
    phenomenon="Maya has $20 to plan her birthday party. She wants a big cake, balloons, games, "
               "AND party favors — but she can't afford everything. What should she choose?",
    driving_question="How do we decide what to buy when we can't buy everything?",
    standards=[
        "SS.1: Explain the difference between needs and wants",
        "SS.2: Identify that choices have consequences (opportunity cost)",
        "FDM.1: Recognize that people must make choices because resources are limited",
    ],
    ldw_phase="LIFE",
    sim_description=(
        "Students help Maya plan her birthday party with a $20 budget. A visual \"money jar\" shows "
        "their remaining funds as they navigate a party supply store. Each item has a price tag. "
        "Students tap to add items to the cart, and the jar drains in real-time. When they overspend, "
        "Mr. Larry gently asks them to put something back. Three different party outcomes result from "
        "their choices — all are fun, teaching that good decisions aren't about having everything."
    ),
    branches=[
        ["Cake or Cupcakes?", "Big cake ($12)", "Cupcakes ($6)", "Make your own ($3 + baking activity)", "Affects remaining budget for other items"],
        ["Decorations", "Balloon arch ($8)", "Regular balloons ($3)", "Paper streamers ($2)", "Visual party look changes based on choice"],
        ["Party Activity", "Bounce house rental ($15 — over budget!)", "Musical chairs (free)", "Treasure hunt ($4 for prizes)", "Teaches that free things can be just as fun"],
        ["Party Favors", "Store-bought bags ($6)", "Homemade bookmarks ($1)", "No favors, more food instead ($0)", "Friends react positively to ALL choices"],
    ],
    family_activity=(
        "Family Budget Challenge: Parent gives child a real or play $10 budget. Together, they plan a "
        "family movie night — choosing snacks, drinks, and a rental movie within budget. Parent Portal "
        "provides a printable planning sheet. Family submits a photo of their movie night setup."
    ),
    assessment=(
        "Students review their party plan, identify what they gave up (opportunity cost in kid terms: "
        "\"What did you say no to?\"), and draw their favorite party moment. Quick 5-question picture-based "
        "quiz on wants vs. needs."
    ),
    parent_portal="Receives the \"Family Movie Night Budget\" activity. Sees child's simulation choices and score. Gets conversation starters: \"What was the hardest thing to say no to?\"",
    child_portal="Runs the Birthday Party simulation. Earns a \"Budget Planner\" badge upon completion. Party result saved as a shareable image.",
    phase_color=LIFE_COLOR,
)

# K-2 Sim 2
add_sim_lesson(
    num="K2-02",
    title="The Lemonade Stand",
    grade="K-2",
    phenomenon="Jayden and his sister want to buy a new soccer ball ($15). Mom says they can earn the money "
               "by running a lemonade stand on Saturday. But wait — they need to buy lemons, cups, and sugar first! "
               "How much will it cost to MAKE money?",
    driving_question="Why do you need to spend money before you can earn money?",
    standards=[
        "EI.1: Identify ways people earn income",
        "SS.3: Explain that saving is choosing not to spend now to use later",
        "FDM.2: Recognize that businesses earn revenue and have costs",
    ],
    ldw_phase="DREAMS",
    sim_description=(
        "Students help Jayden set up a lemonade stand. They start with a $5 loan from Mom (introducing "
        "borrowing at the simplest level). They buy supplies, set a price per cup, choose a location "
        "(front yard vs. park vs. school), and manage the stand through the day. A visual customer counter "
        "tracks sales. Weather events (sunny = more customers, rain = fewer) add unpredictability. "
        "At the end, they count revenue, subtract costs, and discover profit — and whether they earned enough "
        "for the soccer ball."
    ),
    branches=[
        ["Supplies: Lemons", "Real lemons ($3)", "Powdered mix ($1.50)", "—", "Real lemons = premium product, higher customer satisfaction"],
        ["Price Per Cup", "$0.25 (cheap, many buyers)", "$0.50 (moderate)", "$1.00 (expensive, fewer buyers)", "Introduces price-demand relationship visually"],
        ["Location", "Front yard (few walkers)", "Park (moderate traffic)", "School event (high traffic, $2 table fee)", "Location affects customer volume"],
        ["Unexpected: Rain starts!", "Pack up and go home", "Move under a tree, keep selling", "Offer \"rainy day\" half-price sale", "Teaches resilience and problem-solving"],
    ],
    family_activity=(
        "Real Lemonade Stand (or Bake Sale): Families set up an actual small business for a day. "
        "Parent helps child track costs, revenue, and profit using a simple worksheet. Child deposits "
        "\"profit\" into their app Savings Account. Takes photo for family journal."
    ),
    assessment=(
        "Students answer: \"How much did it cost to start? How much did you earn? What's left?\" "
        "Drawing activity: \"Design your own business — what would you sell?\" "
        "3-question quiz matching words: revenue, cost, profit."
    ),
    parent_portal="Receives the \"Family Business Day\" guide with printable cost/revenue tracker. Sees child's simulation profit/loss. Suggested follow-up: open a real savings jar.",
    child_portal="Runs the Lemonade Stand sim. Virtual earnings added to Savings Account. Unlocks \"Young Entrepreneur\" badge.",
    phase_color=DREAMS_COLOR,
)

# K-2 Sim 3
add_sim_lesson(
    num="K2-03",
    title="The Coin Quest",
    grade="K-2",
    phenomenon="Grandma gives Aiden a jar of mixed coins for his birthday. He shakes it — it sounds like a lot! "
               "But when he opens it, some coins are big and some are tiny. The big ones aren't always worth more. "
               "How much money does he really have?",
    driving_question="Why isn't a bigger coin always worth more money?",
    standards=[
        "SS.1: Identify and count coins and bills",
        "SS.4: Compare the value of different coin combinations",
        "FDM.1: Demonstrate the ability to make basic financial decisions",
    ],
    ldw_phase="LIFE",
    sim_description=(
        "Students help Aiden sort and count his coin jar. Interactive coins can be dragged into sorting trays "
        "(pennies, nickels, dimes, quarters). A running total updates as coins are sorted. The twist: "
        "students discover a dime (small) is worth MORE than a nickel (big), anchoring the phenomenon. "
        "After counting, students \"visit\" three stores where items cost different amounts and they must "
        "make exact change. The final store has a toy Aiden wants — can he afford it with what's in the jar?"
    ),
    branches=[
        ["Sorting Strategy", "Sort by size", "Sort by color", "Sort by value", "Value sorting is most efficient — teaches categorization"],
        ["Store 1: Sticker ($0.35)", "Pay with quarters + dime", "Pay with many nickels", "Pay with 35 pennies", "All correct — shows multiple paths to same answer"],
        ["Store 2: Bookmark ($0.75)", "Three quarters", "Seven dimes + nickel", "Mix of coins", "Practices multiple coin combinations"],
        ["Store 3: Toy ($2.50)", "Can afford — buy it!", "Almost enough — save for next time", "Buy a cheaper toy + save the rest", "Introduces saving vs. spending decision"],
    ],
    family_activity=(
        "Coin Jar Count: Family collects loose change around the house for a week. At week's end, "
        "child sorts and counts with parent. Together they decide: spend it, save it, or share it? "
        "Take a photo of the sorted coins."
    ),
    assessment=(
        "Interactive coin-counting quiz (6 questions). Students also draw \"three things I could do with $5\" "
        "showing spend, save, and share options."
    ),
    parent_portal="Receives printable coin identification cards and the week-long Coin Jar Challenge. Sees child's counting accuracy score.",
    child_portal="Coin Quest simulation with drag-and-drop counting. Earns \"Money Counter\" badge. Coin jar total added to virtual savings.",
    phase_color=LIFE_COLOR,
)

# K-2 Sim 4
add_sim_lesson(
    num="K2-04",
    title="The Three Piggy Banks",
    grade="K-2",
    phenomenon="Sophia gets $3 allowance every week. Her mom gives her three piggy banks labeled \"SAVE,\" \"SPEND,\" "
               "and \"SHARE.\" She has to split her money into all three every week. After 4 weeks, she has $12 total — "
               "but it's split up in different ways depending on her choices. Which piggy bank strategy helps her reach "
               "her dream toy?",
    driving_question="What happens to your money when you split it into Save, Spend, and Share?",
    standards=[
        "SS.3: Describe how saving helps people buy things in the future",
        "SS.5: Explain the concept of allocation (dividing resources)",
        "FDM.3: Set a simple financial goal and create a plan to reach it",
    ],
    ldw_phase="DREAMS",
    sim_description=(
        "Students manage Sophia's weekly $3 allowance across three animated piggy banks over 4 simulated weeks. "
        "Each week, they drag dollar bills/coins into SAVE, SPEND, and SHARE banks. The SPEND bank has a \"store\" "
        "where small treats can be bought each week. The SHARE bank fills up and periodically triggers a giving event "
        "(donating to the school book drive). The SAVE bank grows toward Sophia's dream toy ($8 art set). "
        "Students discover the tradeoff: spending more now means waiting longer for the dream toy."
    ),
    branches=[
        ["Week 1: $3 split", "$2 Save, $0.50 Spend, $0.50 Share", "$1 each", "$0.50 Save, $2 Spend, $0.50 Share", "SAVE balance after Week 1 varies: $2, $1, or $0.50"],
        ["Week 2: Temptation!", "Stick to plan", "Best friend has candy — move $1 from Save to Spend", "Give extra to Share (class food drive)", "Tests discipline — consequences visible in piggy bank totals"],
        ["Week 3: Surprise!", "Grandma sends $2 — where does it go?", "Put it ALL in Save (reach goal faster)", "Split it evenly", "Windfall decision — tests whether goals change behavior"],
        ["Week 4: The Result", "Enough saved — buy the art set!", "Almost there — 1 more week needed", "Spent too much — need 3 more weeks", "Visual celebration for savers; gentle encouragement for spenders"],
    ],
    family_activity=(
        "Three Jar System: Family sets up 3 real jars (or envelopes) labeled Save, Spend, Share. "
        "Child allocates their real allowance or chore earnings each week. Parent photographs the jars weekly. "
        "After 4 weeks, family celebrates the SAVE achievement together."
    ),
    assessment=(
        "Students present their 4-week piggy bank results. Reflection questions: \"What was hardest — saving, "
        "spending wisely, or sharing?\" Drawing: \"My dream and how I'll save for it.\" "
        "3-question quiz on saving vs. spending."
    ),
    parent_portal="Receives the Three Jar Setup guide. Weekly prompt to photograph jar progress. Dashboard shows child's simulation allocation pattern over 4 weeks.",
    child_portal="Animated piggy bank simulation. SAVE progress bar tracks toward dream toy. Earns \"Smart Saver\" badge. Virtual allowance system activated.",
    phase_color=DREAMS_COLOR,
)


# ============================================================
# GRADES 3-5 SIMULATIONS
# ============================================================

add_heading("Grade Band: 3-5 (Ages 8-11)", 1)
add_body(
    "Upper elementary students are ready for multi-step financial scenarios involving real math, "
    "comparison shopping, budgeting with constraints, and introductory economic concepts. "
    "Simulations involve 3-way branching, calculators, and visual data dashboards."
)

doc.add_paragraph()

# 3-5 Sim 1
add_sim_lesson(
    num="35-01",
    title="HealthMart Foods: The Grocery Challenge",
    grade="3-5",
    phenomenon="Mom sends you to the grocery store with a list and $40. You need to buy green beans, chicken, "
               "salad, bread, a drink, fruit, milk, and toilet paper. But every item has 3 versions at different prices — "
               "and the cheapest isn't always the best deal. How do you feed your family well without overspending?",
    driving_question="How do you make smart choices when every option has a different price, quality, and trade-off?",
    standards=[
        "SS.2: Apply comparison shopping strategies",
        "SS.6: Calculate unit price and evaluate value",
        "FDM.2: Analyze trade-offs in financial decisions",
        "FDM.4: Demonstrate how to use a budget to manage spending",
    ],
    ldw_phase="LIFE",
    sim_description=(
        "This is the FLAGSHIP simulation — evolved from Mr. Larry's original nutrition module. Students navigate "
        "a virtual HealthMart Foods grocery store. Each aisle presents a product category with 3 options varying "
        "in price, prep time, and health score. A running budget tracker, time tracker, and health meter appear "
        "at the top of the screen. Students must balance all three. Hidden \"coupons\" reward exploration. "
        "A \"sale\" event mid-simulation tests impulse buying. The checkout reveals total cost, health grade, "
        "and time investment — with Mr. Larry providing commentary on the choices made."
    ),
    branches=[
        ["Green Beans", "Fresh ($2.50, 15 min prep, A+)", "Frozen ($1.75, 5 min, A)", "Canned ($1.00, 3 min, B)", "Health vs. price vs. time tradeoff"],
        ["Chicken", "Whole chicken ($5, 1 hr prep, A+)", "Chicken breast ($7, 1 hr, A+)", "Pre-made rotisserie ($9, 0 min, A)", "Convenience premium introduced"],
        ["Bread", "Wheat ($3.50, 3 min, A)", "White ($2.50, 3 min, B)", "Dinner rolls ($2, 0 min, B-)", "Health labels and ingredient reading"],
        ["SURPRISE: Soda BOGO Sale!", "Buy the deal ($4 for 2)", "Skip it — buy water ($1)", "Buy juice instead ($2.50)", "Impulse buying vs. sticking to list"],
        ["Milk", "Non-fat ($3.50, A+)", "Low-fat ($3.25, A)", "Whole ($3.00, B+)", "Nutritional comparison + math"],
        ["Checkout", "Under budget + healthy", "Under budget + unhealthy", "Over budget — must remove items", "Consequence reflects all prior choices"],
    ],
    family_activity=(
        "Real Grocery Trip: Parent gives child $40 (real or budgeted). Child makes the grocery list, "
        "comparison shops at a real store, and tracks spending on a worksheet. Family cooks a meal from "
        "the purchased items together. Discussion: What surprised you about prices?"
    ),
    assessment=(
        "Students present their grocery receipt (simulated) with analysis: total cost, health grade, "
        "time investment. Reflection: \"What would you change next time?\" Unit price calculation worksheet. "
        "5-question quiz on comparison shopping and opportunity cost."
    ),
    parent_portal="Receives the Real Grocery Trip guide with budget worksheet. Sees child's simulation results: budget remaining, health grade, time score. Suggested dinner recipe from child's choices.",
    child_portal="Full HealthMart Foods simulation with real-time trackers. Unlocks \"Smart Shopper\" badge. Can replay for better scores. Leaderboard option for classrooms.",
    phase_color=LIFE_COLOR,
)

# 3-5 Sim 2
add_sim_lesson(
    num="35-02",
    title="The Class Trip Fund",
    grade="3-5",
    phenomenon="Your class wants to go on a field trip to the science museum. It costs $15 per student and there are "
               "25 students. That's $375 total! The school will pay half, but the class needs to raise $187.50. "
               "You have 8 weeks. How do you raise the money?",
    driving_question="How do groups work together to reach a big financial goal?",
    standards=[
        "SS.3: Develop a savings plan to achieve a goal",
        "EI.2: Identify ways to earn money through work and enterprise",
        "FDM.3: Create and evaluate a financial plan",
        "FDM.5: Work collaboratively on financial decisions",
    ],
    ldw_phase="DREAMS",
    sim_description=(
        "Students manage a class fundraising campaign over 8 simulated weeks (compressed into the simulation). "
        "Each week they choose a fundraising strategy (bake sale, car wash, read-a-thon, etc.), allocate "
        "volunteer time, and deal with unexpected events (rain cancels the car wash, a parent donates $50, "
        "the bake sale ingredient cost is higher than expected). A thermometer-style progress bar tracks "
        "funds toward $187.50. Students learn about revenue vs. expenses, teamwork, and contingency planning."
    ),
    branches=[
        ["Week 1-2: Strategy", "Bake sale ($80 revenue, $25 cost)", "Car wash ($60 revenue, $10 cost)", "Read-a-thon (pledges: $50-120, no cost)", "Each has different risk/reward profile"],
        ["Week 3: Setback!", "Reschedule to next week", "Switch to indoor alternative", "Move to online pledge drive", "Teaches resilience and flexibility"],
        ['Week 5: Windfall', 'Parent donates $50 — Put it all toward trip', 'Use $20 to fund a second fundraiser', 'Save it as emergency buffer', 'Windfall management'],
        ['Week 7: Almost there!', 'Short by $30 — One big push (extra bake sale)', 'Ask each student for $1.50', 'Write to local business for sponsorship', 'Multiple solutions to same problem'],
    ],
    family_activity=(
        "Family Fundraising Brainstorm: Together, identify a family financial goal (new game, day trip, etc.). "
        "Create a plan to save toward it in 4 weeks using a thermometer chart on the fridge. "
        "Child proposes ways to earn extra (chores, lemonade stand, etc.)."
    ),
    assessment=(
        "Students present their fundraising plan with a budget breakdown: revenue, costs, profit per event. "
        "Reflection: \"What was your most successful strategy? What would you do differently?\" "
        "Math worksheet: calculating percentages of goal reached."
    ),
    parent_portal="Receives the Family Goal Thermometer template. Sees child's fundraising simulation results. Weekly prompt: \"How much closer are you to your family goal?\"",
    child_portal="Fundraising simulation with week-by-week decisions. Progress thermometer. Earns \"Goal Getter\" badge. Can share results with classmates.",
    phase_color=DREAMS_COLOR,
)

# 3-5 Sim 3
add_sim_lesson(
    num="35-03",
    title="Paycheck Payday",
    grade="3-5",
    phenomenon="Marcus just got his first \"paycheck\" from his weekly chores: $25. But wait — Mom says he owes "
               "$2.50 for the snack she bought him at the store (a \"tax\"), and $1.00 goes to the family emergency jar. "
               "His paycheck was $25, but he only has $21.50 to keep. Why do people get less money than they earn?",
    driving_question="Why is the money you take home different from the money you earned?",
    standards=[
        "EI.1: Explain the difference between gross and net income",
        "EI.3: Identify common paycheck deductions (taxes, benefits)",
        "SS.7: Create and follow a personal budget based on net income",
    ],
    ldw_phase="LIFE",
    sim_description=(
        "Students receive a simulated weekly paycheck from their app-based chore system. The simulation walks "
        "them through the anatomy of a paycheck: gross pay, deductions (simplified as \"community tax\" and "
        "\"family fund\"), and net pay. Then students must budget their net pay across categories: save, spend, "
        "share, and a new category — \"bills\" (they owe for their share of the family streaming service: $2/week). "
        "A visual paycheck stub is generated that they can print and keep."
    ),
    branches=[
        ["Gross Pay: $25", "Accept and continue", "Ask: Why is this amount?", "—", "Curious students get a mini-lesson on how pay is determined"],
        ['Deductions Revealed', '$2.50 community tax + $1 family fund — Student reacts: unfair!', 'Student reacts: I understand', '—', 'Mr. Larry explains: taxes fund shared resources'],
        ["Net Pay: $21.50", "Budget it carefully", "Spend it all at once", "Save everything", "Different outcomes shown after 4 weeks of each strategy"],
        ["Bill Arrives: Streaming $2", "Pay it from Spend", "Oops — already spent everything", "Had it planned — pays easily", "Teaches fixed expenses must come first"],
    ],
    family_activity=(
        "My First Real Budget: Parent shares a simplified version of a real household bill (electric, groceries, "
        "or streaming). Child creates a budget for their allowance that includes a \"bill\" payment. "
        "Discussion: What are our family's biggest expenses?"
    ),
    assessment=(
        "Students annotate a blank paycheck stub with correct labels. Budget creation exercise: "
        "given $20 net pay, allocate across Save, Spend, Share, and Bills. "
        "5-question quiz on gross vs. net, deductions, and budgeting."
    ),
    parent_portal="Receives age-appropriate bill explanation templates. Option to set up a real \"family bill\" in the chore payroll system (child pays $1-2/week toward a family goal). Sees child's budget allocation.",
    child_portal="Paycheck simulator with animated paycheck generation. Budget planner tool. Earns \"Budget Boss\" badge. Paycheck stub saved to portfolio.",
    phase_color=LIFE_COLOR,
)

# 3-5 Sim 4
add_sim_lesson(
    num="35-04",
    title="The Savings Snowball",
    grade="3-5",
    phenomenon="Twin sisters each save $5 per week. Amara puts her money in a jar. Zuri puts hers in a savings "
               "account that earns 2% interest every month. After 6 months, Amara has $130. Zuri has $134.50. "
               "Zuri got FREE money she never earned. How?",
    driving_question="How does money grow when you're not even touching it?",
    standards=[
        "INV.1: Define interest and explain how it affects savings",
        "SS.3: Compare saving strategies (jar vs. bank vs. investment)",
        "FDM.6: Use math to project financial outcomes over time",
    ],
    ldw_phase="WEALTH",
    sim_description=(
        "Students run a side-by-side savings experiment. They set a weekly savings amount and watch two "
        "animated accounts grow over 6 simulated months — one as a plain jar (no interest) and one as a "
        "savings account (2% monthly). An animated \"interest fairy\" visits the savings account each month, "
        "visually adding bonus coins. Students can adjust their weekly savings amount and see how the gap "
        "between jar and account widens over time. A time-lapse slider shows 1 year, 5 years, 10 years."
    ),
    branches=[
        ["Starting Savings Rate", "$3/week (conservative)", "$5/week (moderate)", "$10/week (ambitious)", "All three paths show interest effect at different scales"],
        ['Month 3: Emergency!', 'Need $15 for school supplies — Withdraw from jar', 'Withdraw from savings', "Don't withdraw — ask parent for help", 'Shows impact of withdrawals on compound growth'],
        ['Month 4: Interest Rate Offer', 'Bank offers 3% if you save $10/week — Switch to higher rate', 'Stay at current rate', '—', 'Introduces incentives for saving more'],
        ['Month 6: Results', 'Compare jar vs. savings — Calculator shows 1-year projection', 'Calculator shows 10-year projection', '—', 'Compound interest visual makes the gap dramatic'],
    ],
    family_activity=(
        "Family Savings Experiment: Open a real savings account (or use the app's virtual account). "
        "Set a family goal and contribute weekly for 4 weeks. Track the interest earned (even if small). "
        "Discuss: \"Where else does compound growth happen in life?\""
    ),
    assessment=(
        "Students calculate simple interest on $100 at 5% for 1 year. Comparison chart: jar vs. savings "
        "vs. investment. Reflection: \"Why did Zuri end up with more?\" "
        "Create a personal savings plan with projected growth."
    ),
    parent_portal="Guide to opening a youth savings account. Compound interest calculator. Sees child's simulation results and projected savings growth.",
    child_portal="Dual savings simulator with time-lapse. Virtual savings account now earns simulated interest. Earns \"Interest Expert\" badge.",
    phase_color=WEALTH_COLOR,
)

# 3-5 Sim 5
add_sim_lesson(
    num="35-05",
    title="The Neighborhood Business",
    grade="3-5",
    phenomenon="Three kids on the same block all start businesses: a dog-walking service, a lawn care service, "
               "and a tutoring service. They all charge $10 per job. After one month, the dog walker made $120, "
               "the lawn care kid made $80 but spent $30 on gas, and the tutor made $60 but had zero expenses. "
               "Who actually made the most PROFIT?",
    driving_question="Why is earning the most money not the same as keeping the most money?",
    standards=[
        "EI.2: Distinguish between revenue and profit",
        "EI.4: Identify business expenses and calculate net income",
        "FDM.2: Evaluate opportunity cost in business decisions",
    ],
    ldw_phase="DREAMS",
    sim_description=(
        "Students choose one of three neighborhood businesses and manage it for 4 simulated weeks. "
        "Each business has different revenue potential, expenses, time commitment, and customer demand. "
        "Students set prices, manage supplies, deal with competition (what happens when two dog walkers "
        "are on the same block?), and track profit. A visual P&L (profit and loss) statement builds "
        "week by week. Students discover that the highest revenue doesn't always mean the highest profit."
    ),
    branches=[
        ["Choose Your Business", "Dog walking ($0 startup)", "Lawn care ($30 startup for gas/tools)", "Tutoring ($0 startup, needs skills)", "Each has unique economics"],
        ["Pricing Strategy", "Undercut competition ($8)", "Match the market ($10)", "Premium pricing ($15, fewer clients)", "Price affects volume and total revenue"],
        ['Week 2: Competition!', 'New kid starts same business — Lower prices to compete', 'Offer better service', 'Find a different neighborhood', 'Competition response strategies'],
        ['Week 3: Growth Decision', 'Stay solo — Hire a friend (split profits)', 'Expand to new neighborhood', '—', 'Scaling decisions and partnership economics'],
    ],
    family_activity=(
        "Business Plan Night: Child creates a 1-page business plan for a real service they could offer. "
        "Family reviews it together: What are the costs? What's the revenue? What's the profit? "
        "Optional: child runs the business for real for one week."
    ),
    assessment=(
        "Students present their business P&L after 4 simulated weeks. Calculate: revenue, expenses, profit, "
        "profit margin. Comparison: \"Which business would you choose now and why?\" "
        "5-question quiz on revenue vs. profit, expenses, and pricing."
    ),
    parent_portal="Receives the Family Business Plan template. Sees child's business simulation P&L. Suggested activity: walk through a real small business together.",
    child_portal="Business management simulation with P&L dashboard. Virtual earnings added to savings. Earns \"Young CEO\" badge.",
    phase_color=DREAMS_COLOR,
)

# 3-5 Sim 6
add_sim_lesson(
    num="35-06",
    title="The Community Resource Map",
    grade="3-5",
    phenomenon="Two families in the same city have the same income ($4,000/month). Family A lives near a credit union, "
               "a free tax prep center, and a community garden. Family B lives in a neighborhood with only check-cashing "
               "stores and payday lenders. After one year, Family A has saved $2,400. Family B has saved $200. Same income — "
               "very different results. Why?",
    driving_question="How does where you live affect how much money you can save?",
    standards=[
        "FDM.4: Identify financial resources available in a community",
        "CD.1: Distinguish between helpful and harmful financial services",
        "SS.8: Explain how access to resources affects financial outcomes",
    ],
    ldw_phase="WEALTH",
    sim_description=(
        "Students navigate two side-by-side virtual neighborhoods on a map. They manage monthly finances "
        "for both families — same income, same expenses — but with different available resources. Family A "
        "can use a credit union (low fees), free tax prep, community garden (lower food costs), and a "
        "library (free internet). Family B must use check-cashing stores (high fees), paid tax prep, "
        "and higher food costs. Students watch the savings gap grow month by month and identify the "
        "\"hidden costs\" of living in an under-resourced neighborhood."
    ),
    branches=[
        ["Bank vs. Check Cashing", "Credit union (free checking)", "Check-cashing store ($5/check fee)", "—", "Family B loses $20-40/month on fees alone"],
        ["Tax Season", "Free VITA tax prep (full refund)", "Paid prep ($200 fee + RAL loan)", "—", "Family B loses hundreds; refund anticipation loan costs more"],
        ["Groceries", "Community garden + store ($250/mo)", "Corner store only ($400/mo)", "—", "Food desert effect visible in budget"],
        ["Emergency: Car breaks down", "Credit union loan (5% APR)", "Payday lender (400% APR!)", "No loan — can't fix car, loses job", "Predatory lending trap demonstrated"],
    ],
    family_activity=(
        "Community Resource Hunt: Family walks or drives through their neighborhood and maps free/low-cost "
        "financial resources (library, credit union, community center, food bank). Create a \"Family Resource Map\" "
        "poster. Discussion: What resources do we use? What's missing?"
    ),
    assessment=(
        "Students write a letter to a city official about what financial resources their neighborhood needs. "
        "Comparison chart: credit union vs. payday lender costs over 1 year. "
        "Reflection: \"What would you change about Family B's neighborhood?\""
    ),
    parent_portal="Receives the Community Resource Map activity guide. List of local free financial resources (VITA, credit unions, etc.). Sees child's simulation results.",
    child_portal="Dual-neighborhood simulation with monthly budget tracker. Map-based exploration. Earns \"Community Champion\" badge.",
    phase_color=WEALTH_COLOR,
)


# ============================================================
# GRADES 6-8 SIMULATIONS
# ============================================================

add_heading("Grade Band: 6-8 (Ages 11-14)", 1)
add_body(
    "Middle school students engage with complex multi-path simulations involving percentages, credit, "
    "investing, entrepreneurship, and systemic financial concepts. Simulations feature spreadsheet-style "
    "tools, real-world data, and collaborative multiplayer elements."
)

doc.add_paragraph()

# 6-8 Sim 1
add_sim_lesson(
    num="68-01",
    title="The Credit Score Mystery",
    grade="6-8",
    phenomenon="Destiny and Marcus both want to buy a used car for $8,000. They go to the same bank. "
               "Destiny gets a loan at 4% interest. Marcus gets offered 18% interest — or no loan at all. "
               "The bank won't tell them why directly. Both have jobs. Both have income. "
               "The difference? A three-digit number neither of them fully understands.",
    driving_question="What is a credit score, how is it calculated, and why does it control so much of your financial life?",
    standards=[
        "CD.1: Explain what a credit score is and what factors affect it",
        "CD.2: Describe the relationship between credit scores and borrowing costs",
        "CD.3: Identify actions that build or damage credit",
        "FDM.7: Evaluate the long-term cost of borrowing decisions",
    ],
    ldw_phase="LIFE",
    sim_description=(
        "Students manage two characters' financial lives over 2 simulated years, making decisions that affect "
        "their credit scores in real-time. A live credit score meter (300-850) responds to every choice. "
        "Students discover the five factors: payment history (35%), amounts owed (30%), credit history length (15%), "
        "new credit (10%), and credit mix (10%). At the end, both characters apply for the same car loan — "
        "and students see how their choices directly determined the interest rate, monthly payment, and total cost."
    ),
    branches=[
        ["Month 1: First Credit Card", "Use it responsibly (small purchases, pay in full)", "Max it out immediately ($500 limit)", "Don't get one (no credit history)", "Each path builds a different credit trajectory"],
        ["Month 6: Bill Payment", "Always pay on time", "Miss one payment (forgot)", "Pay minimum only", "Payment history is 35% of score — biggest factor"],
        ["Month 12: New Phone Offer", "Buy outright ($200)", "Finance it (new credit inquiry)", "Put on credit card", "New credit inquiries temporarily lower score"],
        ["Month 18: Car Loan Application", "High score = low rate", "Medium score = medium rate", "Low score = denied or predatory rate", "Final consequence of all prior decisions"],
        ["Total Car Cost Comparison", "$8,800 at 4% / 3 years", "$10,400 at 12% / 3 years", "$12,800 at 18% / 3 years", "Same car — up to $4,000 difference in total cost"],
    ],
    family_activity=(
        "Credit Score Family Discussion: Parent shares their own credit score experience (optional) or "
        "uses the Credit Score Simulator to show a hypothetical. Together, research one action that "
        "builds credit and one that damages it. Older teens: pull a free credit report together at annualcreditreport.com."
    ),
    assessment=(
        "Students create an infographic explaining the 5 credit score factors. Calculate the total cost "
        "of a $10,000 loan at 5%, 10%, and 20% interest over 5 years. "
        "Reflection: \"What will you do in your first year of adulthood to build good credit?\""
    ),
    parent_portal="Credit score education resources. Family discussion guide. Sees child's simulation credit scores for both characters. Suggested: review a real credit card statement together.",
    child_portal="Credit Score Simulator with live meter. Dual-character management over 2 years. Earns \"Credit Wise\" badge. Credit score knowledge quiz.",
    phase_color=LIFE_COLOR,
)

# 6-8 Sim 2
add_sim_lesson(
    num="68-02",
    title="The Investment Challenge",
    grade="6-8",
    phenomenon="In 1990, two friends each received $1,000 as a graduation gift. Alex put it in a savings account "
               "earning 2% per year. Jordan invested it in a stock market index fund averaging 10% per year. "
               "It's now 2026 — 36 years later. Alex has $2,039. Jordan has $30,912. Same starting amount. "
               "Same year. How is that possible?",
    driving_question="How do small differences in where you put your money create huge differences over time?",
    standards=[
        "INV.1: Explain the concept of compound interest and investment growth",
        "INV.2: Compare risk and return across different investment types",
        "INV.3: Define stocks, bonds, mutual funds, and index funds",
        "FDM.6: Use math to project financial outcomes",
    ],
    ldw_phase="WEALTH",
    sim_description=(
        "Students receive a virtual $1,000 and must allocate it across 4 investment types: savings account (2%), "
        "bonds (4%), index fund (10% avg, volatile), and individual stocks (variable, high risk). A time-lapse "
        "simulator shows growth over 10, 20, and 30 years with animated graphs. The stock market simulation "
        "includes real historical events (2008 crash, 2020 pandemic dip, 2024 AI boom) so students experience "
        "volatility. Those who panic-sell during dips lose money. Those who hold steady recover and grow."
    ),
    branches=[
        ["Initial Allocation", "100% savings (safe)", "50/50 savings and index fund", "100% index fund (aggressive)", "Equal split across all 4"],
        ["Year 5: Market Drops 20%!", "Sell everything (lock in losses)", "Hold steady (wait for recovery)", "Buy more (\"buy the dip\")", "Emotional vs. rational decision-making"],
        ["Year 15: Hot Stock Tip", "Put 50% into one company", "Ignore it, stay diversified", "Research first, invest 10%", "Diversification vs. concentration risk"],
        ["Year 30: Results", "See final portfolio value", "Compare to savings-only path", "Calculate: what if you started 10 years later?", "Time in market > timing the market"],
    ],
    family_activity=(
        "Family Investment Research: Pick 3 companies the family uses daily (Nike, Apple, McDonald's, etc.). "
        "Look up their stock price today. Imagine investing $100 in each one year ago — would you have gained or lost? "
        "Discussion: \"Why does a company's stock price go up or down?\""
    ),
    assessment=(
        "Students present their 30-year portfolio results with graphs. Calculate compound growth: "
        "$100/month at 8% for 30 years. Write a letter to their future 18-year-old self with investment advice. "
        "5-question quiz on compound interest, diversification, and risk vs. return."
    ),
    parent_portal="Investment basics guide for families. Retirement calculator tool. Sees child's portfolio simulation results. Suggested: explore a custodial investment account.",
    child_portal="Full investment simulator with real historical data. Portfolio tracker in Child Portal's Investments section. Earns \"Future Investor\" badge.",
    phase_color=WEALTH_COLOR,
)

# 6-8 Sim 3
add_sim_lesson(
    num="68-03",
    title="The Apartment Hunt",
    grade="6-8",
    phenomenon="Three friends graduate from school and need apartments. Each earns $2,500/month. They find three options: "
               "a nice apartment for $1,500/month (60% of income), a decent one for $750/month (30%), and a basic one for "
               "$500/month (20%). The friend who chose the $1,500 apartment looks like she's living the best life on social "
               "media — but she can't afford groceries by the end of the month. What went wrong?",
    driving_question="How much of your income should go to housing, and what happens when you get that ratio wrong?",
    standards=[
        "SS.7: Apply the 50/30/20 budgeting rule",
        "SS.9: Explain the concept of fixed vs. variable expenses",
        "FDM.4: Create a realistic monthly budget based on income",
        "RM.1: Identify the importance of an emergency fund",
    ],
    ldw_phase="LIFE",
    sim_description=(
        "Students choose an apartment and then manage a full monthly budget on $2,500 income. Fixed expenses "
        "(rent, utilities, phone, insurance) are set first. Then variable expenses (groceries, transportation, "
        "entertainment) must be managed week by week. Random life events (car trouble, medical bill, friend's "
        "birthday dinner) test their emergency buffer. Students who chose the expensive apartment face impossible "
        "tradeoffs. A \"social media vs. reality\" meter shows the disconnect between online appearances and "
        "actual financial stress."
    ),
    branches=[
        ["Apartment Choice", "Luxury: $1,500 (60% income)", "Moderate: $750 (30%)", "Basic: $500 (20%)", "Sets the entire budget constraint for the simulation"],
        ["Week 1: Fixed Bills Due", "Rent + utilities + phone = ???", "Luxury: $1,750 in bills", "Moderate: $1,000 in bills", "Basic: $750 in bills — most budget flexibility"],
        ["Week 2: Grocery Budget", "$300/month (healthy)", "$150/month (tight)", "$50/month (ramen only)", "Luxury renter forced into cheapest food option"],
        ["Week 3: Car Breaks Down ($400)", "Pay from savings", "Put on credit card", "Can't afford either — lose transportation", "Emergency fund (or lack thereof) determines outcome"],
        ["End of Month", "Saved money + comfortable", "Broke even — barely", "In debt — credit card growing", "Housing choice cascades through entire financial life"],
    ],
    family_activity=(
        "Family Budget Night: Parent shares the household's approximate monthly budget categories (simplified). "
        "Child calculates what percentage goes to housing, food, transportation, etc. "
        "Discussion: \"Does our family follow the 50/30/20 rule? What would we change?\""
    ),
    assessment=(
        "Students create three complete monthly budgets for the three apartment choices. Write a social media "
        "post that's honest about finances vs. one that only shows the highlights. "
        "Reflection: \"Why do people choose the expensive apartment even when they can't afford it?\""
    ),
    parent_portal="Family budgeting worksheet. Guide to discussing housing costs with kids. Sees child's simulation budget outcomes.",
    child_portal="Full monthly budget simulator with week-by-week events. Social media reality checker. Earns \"Budget Master\" badge.",
    phase_color=LIFE_COLOR,
)

# 6-8 Sim 4
add_sim_lesson(
    num="68-04",
    title="Shark Tank Jr.: The Pitch",
    grade="6-8",
    phenomenon="A 13-year-old invented a phone case that doubles as a fidget toy. She made 50 of them for $3 each "
               "and sold them at school for $12 each. She made $450 profit in two weeks. Now she wants to scale — "
               "but a factory requires a minimum order of 1,000 units at $2 each. That's $2,000 she doesn't have. "
               "Should she take a loan, find an investor, save up, or crowdfund?",
    driving_question="How do entrepreneurs fund their growth, and what does each funding source cost you?",
    standards=[
        "EI.5: Explain different sources of business funding",
        "EI.6: Calculate profit margins and break-even points",
        "CD.4: Compare the cost of different types of borrowing",
        "FDM.8: Evaluate risk vs. reward in business decisions",
    ],
    ldw_phase="DREAMS",
    sim_description=(
        "Students step into the shoes of a teen entrepreneur scaling their business. They choose a product, "
        "set a price, calculate unit economics, and then face the growth dilemma: how to fund 1,000 units. "
        "Four funding paths are simulated: bank loan (pay interest), investor (give up equity), crowdfunding "
        "(pre-sell, risk of underdelivering), or bootstrap (save slowly, risk missing the market). Each path "
        "plays out over 6 simulated months with real financial consequences. A Shark Tank-style pitch at the end "
        "lets students present their business to virtual investors."
    ),
    branches=[
        ["Funding Choice", "Bank loan ($2K at 10%)", "Investor ($2K for 30% equity)", "Crowdfund (pre-sell 200 units)", "Bootstrap (save $500/month for 4 months)"],
        ["Month 2: Production", "Factory delivers on time", "Factory delayed 3 weeks", "Quality issue — 10% defective", "Each funding path has different leverage to handle problems"],
        ['Month 4: Competition!', 'Bigger company copies your product — Lower price to compete', 'Improve quality to differentiate', 'Pivot to new design', 'Market competition response'],
        ["Month 6: Results", "Revenue, expenses, profit, debt/equity given up", "Loan path: profitable but owe payments", "Investor path: profitable but own less", "Bootstrap: smaller but 100% yours"],
    ],
    family_activity=(
        "Family Shark Tank Night: Each family member pitches a business idea (real or fun). "
        "Family \"invests\" virtual money. Discussion: \"Would you rather own 100% of a small business "
        "or 70% of a big business?\" Create a one-page business plan for the best idea."
    ),
    assessment=(
        "Students present a Shark Tank pitch: product, market, financials, funding ask. "
        "Calculate break-even point and profit margin. Compare total cost of each funding method. "
        "Reflection: \"What surprised you about the cost of getting money to make money?\""
    ),
    parent_portal="Family Shark Tank Night guide. Business plan template. Sees child's simulation business results and pitch.",
    child_portal="Full entrepreneur simulation with product development, funding, and scaling. Pitch recorder. Earns \"Shark Tank Survivor\" badge.",
    phase_color=DREAMS_COLOR,
)

# 6-8 Sim 5
add_sim_lesson(
    num="68-05",
    title="The Insurance Gamble",
    grade="6-8",
    phenomenon="Two families live next door to each other. Both have the same house, same cars, same income. "
               "Family A pays $400/month for insurance (health, auto, home). Family B pays $0 — they cancelled everything "
               "to save money. For 3 years, Family B brags about saving $14,400. Then a tree falls on both houses during "
               "a storm. Family A pays a $1,000 deductible. Family B owes $45,000 in repairs. Who made the smarter bet?",
    driving_question="Is insurance a waste of money — or the most important thing you'll ever pay for?",
    standards=[
        "RM.1: Explain the purpose and basic types of insurance",
        "RM.2: Calculate the cost-benefit of insurance premiums vs. potential losses",
        "FDM.9: Evaluate risk probability and financial consequences",
    ],
    ldw_phase="WEALTH",
    sim_description=(
        "Students manage two households over 5 simulated years — one with full insurance, one without. "
        "Random life events are generated each year: car accidents, health emergencies, weather events, "
        "theft, nothing at all. Students see how insurance absorbs shocks while the uninsured family "
        "faces catastrophic costs. A probability visualizer shows that the QUESTION isn't whether something "
        "will happen — it's WHEN. Students adjust coverage levels and deductibles to find the right balance."
    ),
    branches=[
        ["Insurance Strategy", "Full coverage ($400/mo)", "Basic coverage ($200/mo)", "No insurance ($0/mo)", "Each level protects against different risks"],
        ["Year 1: Nothing happens", "Insured: \"wasted\" $4,800", "Uninsured: saved $4,800", "—", "Short-term view favors no insurance"],
        ["Year 3: Car accident ($8,000)", "Insured: $500 deductible", "Basic: $2,000 deductible", "Uninsured: $8,000 out of pocket", "First major event shifts the math"],
        ["Year 4: Health emergency ($25,000)", "Insured: $3,000 max out of pocket", "Basic: $10,000 max", "Uninsured: $25,000 — debt or bankruptcy", "Catastrophic event changes everything"],
        ["Year 5: Totals", "Insured: paid $24K premiums, saved $28K+ in covered losses", "Uninsured: saved $24K premiums, lost $33K+ in uncovered events", "—", "Net math clearly favors insurance"],
    ],
    family_activity=(
        "Insurance Scavenger Hunt: Family identifies all insurance they currently have (auto, health, renters/home, etc.). "
        "Child calculates total monthly cost. Discussion: \"What would happen if we didn't have health insurance and someone "
        "got sick?\" Research one type of insurance the family doesn't have."
    ),
    assessment=(
        "Students create a comparison chart: insured vs. uninsured over 5 years with actual costs. "
        "Write an argument: \"Should a 22-year-old buy health insurance?\" "
        "Calculate: monthly premium vs. potential loss probability."
    ),
    parent_portal="Family insurance inventory checklist. Guide to discussing risk with kids. Sees child's simulation outcomes.",
    child_portal="Dual-household insurance simulator with random events. Risk probability visualizer. Earns \"Risk Manager\" badge.",
    phase_color=WEALTH_COLOR,
)

# 6-8 Sim 6 (abbreviated — keeping pattern)
add_sim_lesson(
    num="68-06",
    title="The Generational Wealth Gap",
    grade="6-8",
    phenomenon="Two families start in 1960. The Johnson family (White) buys a house for $15,000 using a VA loan. "
               "The Williams family (Black) is denied the same loan due to redlining and must rent. By 2026, the "
               "Johnson family home is worth $450,000. The Williams family has paid $380,000 in rent — and owns nothing. "
               "Same era. Same income. Different rules.",
    driving_question="How did historical policies create today's wealth gap, and what can be done about it?",
    standards=[
        "FDM.10: Analyze how historical and systemic factors affect financial outcomes",
        "INV.4: Explain how real estate builds generational wealth",
        "CD.5: Identify discriminatory lending practices and their effects",
    ],
    ldw_phase="WEALTH",
    sim_description=(
        "Students follow two families across three generations (1960, 1990, 2026). At each era, they make financial "
        "decisions — but the simulation reveals that the OPTIONS available differ based on historical policies "
        "(redlining, GI Bill exclusion, predatory lending). Students see how a single denied home loan in 1960 "
        "cascades into a $450,000 wealth gap by 2026. The simulation is not about blame — it's about understanding "
        "how systems create outcomes, and what tools exist today to close the gap."
    ),
    branches=[
        ["1960: Home Purchase", "VA loan approved (Johnson)", "VA loan denied — redlining (Williams)", "—", "Same application, different outcomes based on policy"],
        ["1990: Children's Generation", "Johnson kids inherit equity, attend college", "Williams kids: no equity, must take loans", "—", "Intergenerational wealth transfer begins"],
        ["2020: Grandchildren", "Johnson grandkids: down payment from family", "Williams grandkids: first-generation homebuyers", "—", "Three-generation compounding effect"],
        ["2026: The Gap", "$450K net worth vs. $35K", "Student explores: what policies could help?", "Community land trusts, down payment assistance, credit building programs", "Solutions-focused conclusion"],
    ],
    family_activity=(
        "Family Wealth Story: Family discusses their own generational financial history (as comfortable). "
        "\"Did grandparents own their home? What financial advantages or barriers did our family face?\" "
        "Research one community resource that helps families build wealth today."
    ),
    assessment=(
        "Students create a timeline infographic showing how policies affected wealth accumulation across generations. "
        "Research paper: identify 3 current programs that address the wealth gap. "
        "Reflection: \"What would you change about the financial system?\""
    ),
    parent_portal="Family financial history discussion guide (sensitive, optional). Community wealth-building resource list. Sees child's simulation exploration.",
    child_portal="Three-generation wealth simulator with historical context. Interactive policy timeline. Earns \"Equity Advocate\" badge.",
    phase_color=WEALTH_COLOR,
)


# ============================================================
# GRADES 9-12 SIMULATIONS
# ============================================================

add_heading("Grade Band: 9-12 (Ages 14-18)", 1)
add_body(
    "High school students engage with full financial modeling — tax preparation, portfolio management, "
    "college/career planning, home buying, and long-term wealth building. Simulations use real data, "
    "real formulas, and prepare students for adult financial decisions they'll face within 1-4 years."
)

doc.add_paragraph()

# 9-12 Sim 1
add_sim_lesson(
    num="912-01",
    title="The Life After Graduation Simulator",
    grade="9-12",
    phenomenon="It's June. You just graduated. You have a job offer for $35,000/year, $12,000 in student loans, "
               "and a $1,500 bank account. You need an apartment by August 1st. First month's rent + deposit "
               "is $2,400. Your first paycheck doesn't arrive until August 15th. Welcome to adulting. How do you survive "
               "the first 90 days?",
    driving_question="What does it actually cost to start your adult life, and how do you survive the gap between graduation and your first paycheck?",
    standards=[
        "SS.7: Create a comprehensive personal budget",
        "EI.7: Calculate take-home pay after taxes and deductions",
        "CD.2: Manage student loan repayment strategies",
        "RM.1: Establish an emergency fund",
        "FDM.4: Make financial decisions under real-world constraints",
    ],
    ldw_phase="LIFE",
    sim_description=(
        "The most realistic simulation in the curriculum. Students live the first 90 days after graduation "
        "day-by-day. They receive their actual take-home pay (after taxes — shock!), must find housing, "
        "set up utilities, buy work clothes, figure out transportation, make their first student loan payment, "
        "and deal with the unexpected. A real checking account interface shows their balance fluctuating. "
        "Overdraft warnings, late payment penalties, and \"adulting surprises\" (parking ticket, work lunch expectations, "
        "friend's wedding gift) make this intensely practical."
    ),
    branches=[
        ["Housing", "Studio apartment ($1,200/mo)", "Roommate ($700/mo each)", "Stay with parents (free, 45-min commute)", "Each choice cascades through transportation, time, and social life"],
        ["Transportation", "Buy a used car ($3,000 — need a loan)", "Public transit ($100/mo pass)", "Bike + occasional rideshare", "Car = freedom + debt; transit = savings + constraints"],
        ["Student Loan Strategy", "Standard repayment ($127/mo)", "Income-driven repayment ($85/mo)", "Defer 6 months (interest accrues)", "Long-term cost differs dramatically"],
        ["Day 45: Unexpected Medical Bill ($800)", "Pay from savings", "Payment plan ($100/mo)", "Put on credit card (18% APR)", "Emergency fund presence determines outcome"],
        ["Day 90: Financial Health Check", "Stable, building savings", "Surviving but tight", "In growing debt", "Reflects all accumulated decisions"],
    ],
    family_activity=(
        "Launch Cost Calculator: Student and parent calculate the real cost of the first 90 days after graduation. "
        "Include: deposit, rent, utilities setup, groceries, work wardrobe, transportation. \"What's the real number?\" "
        "Discussion: \"What did we wish we'd known when we started out?\""
    ),
    assessment=(
        "Students present their complete 90-day financial plan with monthly budgets, cash flow projections, "
        "and emergency contingencies. Calculate true take-home pay from a $35K salary. "
        "Write a \"letter to my 17-year-old self\" with financial advice."
    ),
    parent_portal="Launch cost calculator. \"What I wish I knew\" sharing prompt. Sees child's simulation financial health score.",
    child_portal="Full 90-day life simulator with checking account interface. Real paycheck calculator. Earns \"Adult Ready\" badge.",
    phase_color=LIFE_COLOR,
)

# 9-12 Sim 2
add_sim_lesson(
    num="912-02",
    title="The Tax Return",
    grade="9-12",
    phenomenon="You worked all year and earned $28,000. Your employer took out $3,360 in taxes. "
               "Your coworker earned the same amount but got a $2,100 refund. You only got $400 back. "
               "She has a kid. You don't. She used the EITC. You forgot to claim your student loan interest. "
               "Same job, same pay — very different tax outcomes.",
    driving_question="How do taxes actually work, and how can understanding them put money back in your pocket?",
    standards=[
        "EI.8: Explain how federal and state income taxes are calculated",
        "EI.9: Identify common deductions, credits, and their effects on tax liability",
        "FDM.11: Complete a basic tax return (1040-EZ equivalent)",
    ],
    ldw_phase="LIFE",
    sim_description=(
        "Students complete a simulated tax return step by step. They receive a W-2, enter their income, "
        "learn the difference between deductions and credits, and discover the EITC (Earned Income Tax Credit), "
        "student loan interest deduction, and education credits. The simulation runs three scenarios side by side: "
        "single filer, single parent, married couple — showing how the same income produces different tax outcomes. "
        "Students use a real 1040 form (simplified) and calculate their refund or amount owed."
    ),
    branches=[
        ["Filing Status", "Single", "Head of Household (with dependent)", "Married Filing Jointly", "Each status has different standard deduction and brackets"],
        ["Deductions", "Standard deduction ($14,600)", "Itemize (student loan interest + charitable)", "—", "Most young filers benefit from standard deduction"],
        ["Credits", "Skip (don't know about them)", "Claim EITC ($600 credit)", "Claim education credit ($1,000)", "Credits > deductions — dollar-for-dollar tax reduction"],
        ["Refund or Owe?", "Refund: invest it", "Refund: spend it", "Owe: payment plan", "What you do with a refund matters as much as getting one"],
    ],
    family_activity=(
        "Tax Basics Night: Parent walks through a simplified version of the family's tax situation (or uses "
        "the simulation's example). Child identifies: income sources, deductions, credits. "
        "Discussion: \"Where does tax money go? What services do we benefit from?\""
    ),
    assessment=(
        "Students complete a full simulated 1040 return for a sample scenario. Calculate: marginal vs. effective "
        "tax rate. Compare three filing statuses for the same income. "
        "Reflection: \"What surprised you most about how taxes work?\""
    ),
    parent_portal="Tax basics family guide. Age-appropriate tax conversation starters. Sees child's simulated return results.",
    child_portal="Step-by-step tax return simulator with real 1040 form. Tax calculator tool. Earns \"Tax Pro\" badge.",
    phase_color=LIFE_COLOR,
)

# 9-12 Sim 3
add_sim_lesson(
    num="912-03",
    title="College vs. Career: The $100K Decision",
    grade="9-12",
    phenomenon="Twins Aaliyah and Devon graduate together. Aaliyah goes to a 4-year university ($25K/year). Devon starts "
               "an apprenticeship earning $35K immediately. At age 26, Aaliyah has a degree, $60K in debt, and a $55K job. "
               "Devon has zero debt, 4 years of experience, and earns $52K. At age 40, Aaliyah earns $85K. Devon earns $70K. "
               "At age 65, who has more total wealth?",
    driving_question="Is college always worth the investment, and how do you calculate the real return on education?",
    standards=[
        "EI.10: Evaluate the financial return on different education and career paths",
        "CD.6: Calculate the true cost of student loan debt over time",
        "FDM.12: Make education and career decisions using financial analysis",
        "INV.5: Apply the concept of opportunity cost to major life decisions",
    ],
    ldw_phase="DREAMS",
    sim_description=(
        "Students build a full financial life model for both paths over 40 years. They factor in: "
        "tuition costs, student loan interest, lost income during college, starting salary, salary growth rate, "
        "savings rate, investment returns, and career advancement. The simulation reveals that the answer isn't "
        "always \"go to college\" — it depends on the field, the school cost, the alternative, and the individual. "
        "Students explore 6 career paths: trades, military, community college + transfer, state university, "
        "private university, and entrepreneurship."
    ),
    branches=[
        ["Education Path", "4-year private ($50K/yr)", "4-year state ($15K/yr)", "Community college + transfer ($8K total)", "Trade apprenticeship ($0 + earn)"],
        ["Year 5: Career Launch", "College grad: $50K starting + $60K debt", "Trade: $55K earning + $0 debt", "CC+transfer: $48K + $15K debt", "Each path has different ROI timeline"],
        ["Year 15: Mid-Career", "Degree holders: salary ceiling higher", "Trades: salary plateaus earlier", "Entrepreneurs: high variance", "Long-term earning curves cross at different points"],
        ["Year 40: Retirement", "Total lifetime earnings calculated", "Total debt paid calculated", "Net wealth comparison", "The \"right\" answer depends on the individual path"],
    ],
    family_activity=(
        "Career + Cost Calculator: Student researches their top 3 career interests. For each: "
        "education required, cost of that education, starting salary, 10-year salary. Family discusses: "
        "\"What can we afford? What scholarships are available? What's our plan?\""
    ),
    assessment=(
        "Students present a 40-year financial model for their chosen career path. Calculate total cost of education, "
        "ROI timeline, and projected retirement wealth. Compare at least 2 paths. "
        "Reflection: \"What factors beyond money matter in choosing your path?\""
    ),
    parent_portal="College cost calculator. Scholarship resource database. Family career planning guide. Sees child's career path analysis.",
    child_portal="40-year life financial modeler with career path comparison. Earns \"Future Planner\" badge.",
    phase_color=DREAMS_COLOR,
)

# 9-12 Sim 4
add_sim_lesson(
    num="912-04",
    title="Building Generational Wealth",
    grade="9-12",
    phenomenon="Your great-grandmother starts saving $50/month in 1970. She invests it in an index fund averaging "
               "10%/year. She never touches it. By 2026, that account has grown to over $1.2 million. She passes it to "
               "your grandmother, who adds to it. Your grandmother passes it to your parent. Now it's coming to you. "
               "One woman's $50/month decision 56 years ago created a millionaire family. What if she had started "
               "with just $25/month? What if she started 10 years later?",
    driving_question="How do ordinary families build extraordinary wealth across generations?",
    standards=[
        "INV.1: Apply compound interest to long-term wealth building",
        "INV.6: Create a diversified investment portfolio",
        "FDM.13: Develop a multi-generational financial plan",
        "FDM.14: Understand estate planning basics (wills, trusts, beneficiaries)",
    ],
    ldw_phase="WEALTH",
    sim_description=(
        "The capstone simulation. Students manage a family's wealth across 3 generations. They set savings rates, "
        "choose investment allocations, navigate life events (marriage, kids, job loss, windfalls), and decide "
        "how wealth transfers between generations. The simulation shows how small, consistent actions compound "
        "into transformative wealth — and how one generation's mistakes or misfortune can set the next generation back. "
        "Students graduate with a personal Generational Wealth Blueprint."
    ),
    branches=[
        ["Generation 1: Foundation", "Save aggressively ($200/mo)", "Save moderately ($100/mo)", "Live paycheck to paycheck", "Foundation set for everything that follows"],
        ["Generation 2: Growth", "Continue saving + invest", "Spend inheritance", "Use inheritance as home down payment", "Intergenerational wealth transfer test"],
        ["Generation 3: You", "Inherit wealth + add to it", "Inherit debt (reverse wealth transfer)", "Start from zero but with knowledge", "Every path has a viable strategy"],
        ["Your Decision", "Build for Generation 4", "Enjoy it now", "Give back to community", "Values-driven financial planning"],
    ],
    family_activity=(
        "Family Legacy Letter: Together, write a letter to the next generation about the family's financial values, "
        "lessons learned, and hopes. Start a real family investment plan (even $10/month). "
        "Discussion: \"What do we want to pass on — beyond money?\""
    ),
    assessment=(
        "Students create a Generational Wealth Blueprint: personal savings plan, investment allocation, "
        "estate planning basics, and a legacy vision statement. Present to the class/family. "
        "Final reflection: \"What is wealth to me?\""
    ),
    parent_portal="Family legacy planning guide. Will/trust basics for parents. Generational wealth discussion framework. Sees child's complete blueprint.",
    child_portal="Three-generation wealth simulator. Personal Generational Wealth Blueprint creator. Earns the final \"Wealth Builder\" badge — completing the LDW journey.",
    phase_color=WEALTH_COLOR,
)


# ============================================================
# CLOSING
# ============================================================

add_heading("Summary: Complete Simulation Library", 1)

summary_headers = ["#", "Title", "Grade", "Phase", "Standards"]
summary_rows = [
    ["K2-01", "The Birthday Party Budget", "K-2", "LIFE", "SS, FDM"],
    ["K2-02", "The Lemonade Stand", "K-2", "DREAMS", "EI, SS, FDM"],
    ["K2-03", "The Coin Quest", "K-2", "LIFE", "SS, FDM"],
    ["K2-04", "The Three Piggy Banks", "K-2", "DREAMS", "SS, FDM"],
    ["35-01", "HealthMart Foods: The Grocery Challenge", "3-5", "LIFE", "SS, FDM"],
    ["35-02", "The Class Trip Fund", "3-5", "DREAMS", "SS, EI, FDM"],
    ["35-03", "Paycheck Payday", "3-5", "LIFE", "EI, SS"],
    ["35-04", "The Savings Snowball", "3-5", "WEALTH", "INV, SS, FDM"],
    ["35-05", "The Neighborhood Business", "3-5", "DREAMS", "EI, FDM"],
    ["35-06", "The Community Resource Map", "3-5", "WEALTH", "FDM, CD, SS"],
    ["68-01", "The Credit Score Mystery", "6-8", "LIFE", "CD, FDM"],
    ["68-02", "The Investment Challenge", "6-8", "WEALTH", "INV, FDM"],
    ["68-03", "The Apartment Hunt", "6-8", "LIFE", "SS, FDM, RM"],
    ["68-04", "Shark Tank Jr.: The Pitch", "6-8", "DREAMS", "EI, CD, FDM"],
    ["68-05", "The Insurance Gamble", "6-8", "WEALTH", "RM, FDM"],
    ["68-06", "The Generational Wealth Gap", "6-8", "WEALTH", "FDM, INV, CD"],
    ["912-01", "The Life After Graduation Simulator", "9-12", "LIFE", "SS, EI, CD, RM, FDM"],
    ["912-02", "The Tax Return", "9-12", "LIFE", "EI, FDM"],
    ["912-03", "College vs. Career: The $100K Decision", "9-12", "DREAMS", "EI, CD, FDM, INV"],
    ["912-04", "Building Generational Wealth", "9-12", "WEALTH", "INV, FDM"],
]
add_table(summary_headers, summary_rows, col_widths=[0.7, 2.5, 0.6, 0.8, 1.9])

doc.add_paragraph()
doc.add_paragraph()

closing_p = doc.add_paragraph()
closing_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = closing_p.add_run("Life Dreams Wealth")
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = NAVY
run.font.name = 'Calibri'

tag = doc.add_paragraph()
tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tag.add_run('"Where Families Build Financial Futures Together"')
run.italic = True
run.font.size = Pt(12)
run.font.color.rgb = DARK_GOLD
run.font.name = 'Calibri'

contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = contact.add_run("Alexandria's Design LLC | Dr. Charles Martin | 310-709-4893")
run.font.size = Pt(10)
run.font.color.rgb = MID_GRAY
run.font.name = 'Calibri'

# Save
doc.save(OUT)
print(f"Saved: {OUT}")
print(f"Size: {os.path.getsize(OUT) / 1024:.0f} KB")
